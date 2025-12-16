#!/usr/bin/env python3
"""
Generate Tutorial Dataset: Project Phantom - Maritime Fraud Investigation

Creates 15 realistic investigative documents to demonstrate all ArkhamMirror features.
Based on TUTORIAL_DATASET_DESIGN.md specification.

Usage:
    python scripts/generate_tutorial_dataset.py

Output:
    data/tutorial_case/ directory with 15 documents
"""

import os
import sys
from pathlib import Path
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import email.utils

# Add arkham_mirror to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / "arkham_mirror"))

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image, ImageDraw, ImageFont
except ImportError as e:
    print(f"ERROR: Missing required library: {e}")
    print("\nInstall dependencies:")
    print("pip install reportlab python-docx pillow")
    sys.exit(1)


class TutorialDatasetGenerator:
    """Generates the Project Phantom tutorial case documents."""

    def __init__(self, output_dir: str = "data/tutorial_case"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Base dates (relative to 2024)
        self.date_incorporation = datetime(2024, 1, 15)
        self.date_first_shipment = datetime(2024, 2, 10)
        self.date_whistleblower = datetime(2024, 3, 15)
        self.date_inspection = datetime(2024, 4, 1)
        self.date_wire_transfer = datetime(2024, 4, 15)
        self.date_second_shipment = datetime(2024, 5, 1)
        self.date_final_report = datetime(2024, 5, 15)

        print(f"[*] Output directory: {self.output_dir.absolute()}")

    def generate_all(self):
        """Generate all 15 tutorial documents."""
        print("\n[*] Generating Project Phantom Tutorial Dataset...\n")

        generators = [
            (1, "Company_Formation_Certificate.pdf", self.gen_company_formation_cert),
            (2, "Email_Thread_Shipment_C999.eml", self.gen_email_thread_c999),
            (3, "Shipping_Manifest_C999.docx", self.gen_shipping_manifest),
            (4, "Warehouse_Inventory_Log.txt", self.gen_warehouse_inventory),
            (5, "Whistleblower_Complaint.pdf", self.gen_whistleblower_complaint),
            (6, "Port_Inspection_Report_April1.pdf", self.gen_port_inspection),
            (7, "Bank_Wire_Transfer_Receipt.jpg", self.gen_wire_transfer_image),
            (8, "Handwritten_Note_Meeting.jpg", self.gen_handwritten_note),
            (9, "Email_Chain_Rivera_Chen.eml", self.gen_email_chain_rivera_chen),
            (10, "Customs_Declaration_Form_C999.pdf", self.gen_customs_declaration),
            (11, "Employee_List_Phantom_Logistics.docx", self.gen_employee_list),
            (12, "Phone_Records_Extract.txt", self.gen_phone_records),
            (13, "GPS_Coordinates_Warehouse.txt", self.gen_gps_coordinates),
            (14, "Invoice_Acme_to_Phantom.pdf", self.gen_invoice),
            (15, "Final_Report_Customs_Investigation.docx", self.gen_final_report),
        ]

        for num, filename, generator_func in generators:
            try:
                print(f"[{num}/15] Generating {filename}...")
                generator_func()
                print(f"    [OK] Created: {filename}")
            except Exception as e:
                print(f"    [ERROR] {e}")
                import traceback
                traceback.print_exc()

        # Generate README
        self.gen_readme()
        print(f"\n[OK] All documents generated successfully!")
        print(f"\n[*] Files saved to: {self.output_dir.absolute()}")
        print(f"\n[*] Next steps:")
        print(f"   1. Start ArkhamMirror: reflex run")
        print(f"   2. Navigate to /ingest page")
        print(f"   3. Upload all files from {self.output_dir}")
        print(f"   4. Follow tutorial walkthrough in TUTORIAL_DATASET_DESIGN.md\n")

    # ========== PDF Generators ==========

    def gen_company_formation_cert(self):
        """Generate Company_Formation_Certificate.pdf (scanned-style PDF)."""
        filepath = self.output_dir / "Company_Formation_Certificate.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(width/2, height - inch, "CERTIFICATE OF INCORPORATION")

        # Government seal placeholder
        c.setFont("Helvetica", 10)
        c.drawCentredString(width/2, height - 1.5*inch, "[CAYMAN ISLANDS GOVERNMENT SEAL]")

        # Form content
        c.setFont("Helvetica", 12)
        y = height - 2.5*inch

        lines = [
            f"Certificate Number: CI-2024-001523",
            f"",
            f"Date of Incorporation: {self.date_incorporation.strftime('%B %d, %Y')}",
            f"",
            f"Company Name: Phantom Logistics LLC",
            f"",
            f"Registered Address: 123 Offshore Drive, George Town, Grand Cayman",
            f"Cayman Islands KY1-1234",
            f"",
            f"Nature of Business: Logistics and Shipping Services",
            f"",
            f"Authorized Capital: $10,000,000 USD",
            f"",
            f"Directors and Officers:",
            f"  • Sarah Chen - Chief Executive Officer",
            f"  • David Kim - Warehouse Manager (USA)",
            f"",
            f"Registered Agent: Cayman Corporate Services Ltd.",
            f"",
            f"This certificate is issued pursuant to the Companies Law (2023 Revision)",
            f"of the Cayman Islands.",
            f"",
            f"",
            f"_______________________________",
            f"Registrar of Companies",
            f"Cayman Islands",
        ]

        for line in lines:
            c.drawString(inch, y, line)
            y -= 0.3*inch

        # Footer note (simulating OCR challenge)
        c.setFont("Helvetica-Oblique", 8)
        c.drawString(inch, inch, "CERTIFIED TRUE COPY - Official Use Only")

        c.save()

    def gen_whistleblower_complaint(self):
        """Generate Whistleblower_Complaint.pdf."""
        filepath = self.output_dir / "Whistleblower_Complaint.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        # Header
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width/2, height - inch, "CONFIDENTIAL WHISTLEBLOWER COMPLAINT")

        # Metadata
        c.setFont("Helvetica", 10)
        c.drawString(inch, height - 1.5*inch, f"Filed: {self.date_whistleblower.strftime('%B %d, %Y')}")
        c.drawString(inch, height - 1.7*inch, "To: Port Authority of Los Angeles - Internal Affairs")
        c.drawString(inch, height - 1.9*inch, "From: [ANONYMOUS - Identity Protected]")

        # Body
        c.setFont("Helvetica", 11)
        y = height - 2.5*inch

        content = [
            "Subject: Irregular Conduct by Port Inspector Marcus Rivera",
            "",
            "I am writing to report serious violations of port security protocols and suspected",
            "corruption involving Inspector Marcus Rivera (Badge #4521).",
            "",
            "Summary of Allegations:",
            "",
            "1. BRIBERY: Inspector Rivera has been observed accepting cash payments from",
            "   representatives of Phantom Logistics LLC in exchange for expedited clearance",
            "   of shipping containers.",
            "",
            "2. ILLEGAL ACTIVITY: Containers from Phantom Logistics are bypassing standard",
            "   customs inspections. I have personally witnessed Inspector Rivera waving",
            "   through containers marked 'Priority - No Inspection Required' without proper",
            "   authorization.",
            "",
            "3. DOCUMENT FALSIFICATION: Inspection reports for Container C-999 (arrived",
            f"   {self.date_first_shipment.strftime('%B %d, %Y')}) show 'No Irregularities Found', but I observed",
            "   the container was never opened for physical inspection.",
            "",
            "Evidence Available:",
            "  • Warehouse logs showing irregular clearance times",
            "  • Witness testimony from other warehouse staff",
            "  • Photos of cash exchange (can be provided upon request)",
            "",
            "I request protection under the Port Authority Whistleblower Protection Act.",
            "I am willing to testify if granted anonymity.",
            "",
            "This is a matter of national security. These containers may contain contraband,",
            "smuggled goods, or worse. Immediate investigation is required.",
            "",
            "[Signature Redacted]",
            "Warehouse Manager, Long Beach Facility",
        ]

        for line in content:
            if y < 1.5*inch:
                c.showPage()
                y = height - inch
                c.setFont("Helvetica", 11)
            c.drawString(inch, y, line)
            y -= 0.25*inch

        # Footer
        c.setFont("Helvetica-Bold", 9)
        c.setFillColor(colors.red)
        c.drawString(inch, inch, "CONFIDENTIAL - DO NOT DISTRIBUTE")

        c.save()

    def gen_port_inspection(self):
        """Generate Port_Inspection_Report_April1.pdf."""
        filepath = self.output_dir / "Port_Inspection_Report_April1.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        # Header
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width/2, height - inch, "PORT OF LOS ANGELES")
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(width/2, height - 1.3*inch, "Container Inspection Report")

        # Form fields
        c.setFont("Helvetica", 11)
        y = height - 2*inch

        fields = [
            f"Inspection Date: {self.date_inspection.strftime('%B %d, %Y')}",
            f"Inspector: Marcus Rivera (Badge #4521)",
            f"Location: Pier 9, Terminal C",
            f"",
            f"Container ID: C-999",
            f"Shipper: Global Maritime Services (Hong Kong)",
            f"Consignee: Phantom Logistics LLC",
            f"Declared Contents: Electronics - 500 units",
            f"Declared Weight: 5,000 kg",
            f"Declared Value: $250,000 USD",
            f"",
            f"Inspection Type: ☑ Visual ☐ Physical ☐ X-Ray",
            f"",
            f"Inspection Notes:",
            f"Container seals intact. No visible signs of tampering. Documentation reviewed",
            f"and found to be in order. Manifest matches bill of lading. No discrepancies",
            f"observed.",
            f"",
            f"FINDINGS: ☑ No Irregularities Found ☐ Refer to Customs",
            f"",
            f"Container cleared for release to consignee.",
            f"",
            f"",
            f"_______________________________    ________________",
            f"Inspector Signature                Date",
            f"Marcus Rivera, Badge #4521",
            f"",
            f"",
            f"Supervisor Review: ☑ Approved",
            f"",
        ]

        for line in fields:
            c.drawString(inch, y, line)
            y -= 0.3*inch

        # Official stamp placeholder
        c.setFont("Helvetica-Oblique", 9)
        c.drawString(width - 2.5*inch, 2*inch, "[OFFICIAL PORT AUTHORITY STAMP]")

        c.save()

    def gen_customs_declaration(self):
        """Generate Customs_Declaration_Form_C999.pdf."""
        filepath = self.output_dir / "Customs_Declaration_Form_C999.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        # Header
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width/2, height - inch, "U.S. CUSTOMS AND BORDER PROTECTION")
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(width/2, height - 1.3*inch, "Entry/Immediate Delivery - Form 3461")

        c.setFont("Helvetica", 10)
        y = height - 2*inch

        fields = [
            f"Entry Number: 2024-LA-C999-001",
            f"Date of Entry: {self.date_first_shipment.strftime('%m/%d/%Y')}",
            f"Port of Entry: Los Angeles/Long Beach (2704)",
            f"",
            f"IMPORTER OF RECORD:",
            f"  Name: Phantom Logistics LLC",
            f"  Address: 456 Commerce Blvd, Long Beach, CA 90802",
            f"  Tax ID: 94-1234567",
            f"",
            f"SHIPPER/EXPORTER:",
            f"  Name: Global Maritime Services Ltd.",
            f"  Address: 88 Kowloon Road, Hong Kong",
            f"",
            f"CONTAINER INFORMATION:",
            f"  Container #: C-999",
            f"  Vessel: MV Pacific Trader",
            f"  Voyage: PT-2024-02",
            f"  Country of Origin: China",
            f"",
            f"MERCHANDISE DESCRIPTION:",
            f"  HTS Code: 8542.31.0000",
            f"  Description: Integrated Circuits and Microassemblies",
            f"  Quantity: 500 units",
            f"  Weight: 5,000 kg (11,023 lbs)",
            f"",
            f"DECLARED VALUE:",
            f"  Invoice Value: $250,000.00 USD",
            f"  Insurance: $25,000.00 USD",
            f"  Freight: $15,000.00 USD",
            f"  TOTAL VALUE FOR DUTY: $290,000.00 USD",
            f"",
            f"DUTIES AND FEES:",
            f"  Estimated Duty (0%): $0.00",
            f"  Processing Fee: $500.00",
            f"  Harbor Maintenance Fee: $725.00",
            f"  TOTAL PAYABLE: $1,225.00",
            f"",
            f"DECLARATION:",
            f"I declare that the information provided is true and correct to the best of",
            f"my knowledge.",
            f"",
            f"_______________________________    ________________",
            f"Authorized Signature                Date",
            f"Sarah Chen, CEO - Phantom Logistics LLC",
        ]

        for line in fields:
            c.drawString(inch, y, line)
            y -= 0.25*inch

        c.save()

    def gen_invoice(self):
        """Generate Invoice_Acme_to_Phantom.pdf."""
        filepath = self.output_dir / "Invoice_Acme_to_Phantom.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        # Header
        c.setFont("Helvetica-Bold", 16)
        c.drawString(inch, height - inch, "INVOICE")

        # Company info
        c.setFont("Helvetica-Bold", 12)
        c.drawString(inch, height - 1.5*inch, "Acme Imports Inc.")
        c.setFont("Helvetica", 10)
        c.drawString(inch, height - 1.7*inch, "789 Trade Center Drive")
        c.drawString(inch, height - 1.9*inch, "Los Angeles, CA 90012")
        c.drawString(inch, height - 2.1*inch, "Tax ID: 95-7654321")

        # Invoice details
        c.setFont("Helvetica-Bold", 10)
        c.drawString(width - 2.5*inch, height - 1.5*inch, "Invoice #: ACM-2024-0415")
        c.setFont("Helvetica", 10)
        c.drawString(width - 2.5*inch, height - 1.7*inch, f"Date: {self.date_wire_transfer.strftime('%B %d, %Y')}")
        c.drawString(width - 2.5*inch, height - 1.9*inch, "Due: Upon Receipt")

        # Bill to
        c.setFont("Helvetica-Bold", 11)
        c.drawString(inch, height - 2.8*inch, "BILL TO:")
        c.setFont("Helvetica", 10)
        c.drawString(inch, height - 3*inch, "Phantom Logistics LLC")
        c.drawString(inch, height - 3.2*inch, "456 Commerce Blvd")
        c.drawString(inch, height - 3.4*inch, "Long Beach, CA 90802")

        # Line items table
        y = height - 4.2*inch
        c.setFont("Helvetica-Bold", 10)
        c.drawString(inch, y, "Description")
        c.drawString(width - 3*inch, y, "Quantity")
        c.drawString(width - 1.5*inch, y, "Amount")

        # Draw line
        c.line(inch, y - 0.1*inch, width - inch, y - 0.1*inch)

        # Item
        y -= 0.4*inch
        c.setFont("Helvetica", 10)
        c.drawString(inch, y, "Professional Consulting Services - Q1 2024")
        c.drawString(inch, y - 0.2*inch, "  • Strategic Advisory")
        c.drawString(inch, y - 0.4*inch, "  • Market Analysis")
        c.drawString(inch, y - 0.6*inch, "  • Risk Management")
        c.drawString(width - 3*inch, y, "1")
        c.drawString(width - 1.8*inch, y, "$2,500,000.00")

        # Total
        y -= 1.2*inch
        c.line(width - 3*inch, y, width - inch, y)
        y -= 0.3*inch
        c.setFont("Helvetica-Bold", 12)
        c.drawString(width - 3*inch, y, "TOTAL:")
        c.drawString(width - 1.8*inch, y, "$2,500,000.00")

        # Payment instructions
        y -= 0.8*inch
        c.setFont("Helvetica", 9)
        c.drawString(inch, y, "Wire transfer details:")
        c.drawString(inch, y - 0.2*inch, "Bank: First National Bank of California")
        c.drawString(inch, y - 0.4*inch, "Account: 1234567890")
        c.drawString(inch, y - 0.6*inch, "Routing: 122000661")

        # Footer
        c.setFont("Helvetica-Oblique", 8)
        c.drawCentredString(width/2, inch, "Thank you for your business")

        c.save()

    # ========== DOCX Generators ==========

    def gen_shipping_manifest(self):
        """Generate Shipping_Manifest_C999.docx."""
        filepath = self.output_dir / "Shipping_Manifest_C999.docx"
        doc = Document()

        # Title
        title = doc.add_heading("SHIPPING MANIFEST", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header info
        doc.add_paragraph(f"Manifest #: GM-2024-C999-HK-LA")
        doc.add_paragraph(f"Date: {self.date_first_shipment.strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Vessel: MV Pacific Trader")
        doc.add_paragraph(f"Voyage: PT-2024-02")
        doc.add_paragraph("")

        # Shipper/Consignee
        doc.add_heading("Shipper", 2)
        doc.add_paragraph("Global Maritime Services Ltd.\n88 Kowloon Road\nHong Kong")

        doc.add_heading("Consignee", 2)
        doc.add_paragraph("Phantom Logistics LLC\n456 Commerce Blvd\nLong Beach, CA 90802\nUSA")

        doc.add_paragraph("")

        # Container details
        doc.add_heading("Container Information", 2)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        data = [
            ("Container Number:", "C-999"),
            ("Container Type:", "40ft High Cube"),
            ("Seal Number:", "SL-999-2024"),
            ("Port of Loading:", "Hong Kong"),
            ("Port of Discharge:", "Los Angeles, CA"),
            ("Final Destination:", "Long Beach, CA"),
        ]

        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        doc.add_paragraph("")

        # Cargo details
        doc.add_heading("Cargo Details", 2)
        cargo_table = doc.add_table(rows=2, cols=5)
        cargo_table.style = 'Light Grid Accent 1'

        # Header row
        headers = ["Description", "HS Code", "Quantity", "Weight (kg)", "Value (USD)"]
        for i, header in enumerate(headers):
            cell = cargo_table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data row
        cargo_data = ["Electronics - Integrated Circuits", "8542.31.0000", "500 units", "5,000", "$250,000"]
        for i, data in enumerate(cargo_data):
            cargo_table.rows[1].cells[i].text = data

        doc.add_paragraph("")

        # Notes
        doc.add_heading("Special Instructions", 2)
        doc.add_paragraph("• Handle with care - fragile electronics")
        doc.add_paragraph("• Keep dry - moisture sensitive")
        doc.add_paragraph("• Priority clearance requested")
        doc.add_paragraph("• Contact: Sarah Chen, (310) 555-0123")

        doc.add_paragraph("")
        doc.add_paragraph("_" * 40)
        doc.add_paragraph("Authorized Signature: Sarah Chen, CEO")
        doc.add_paragraph(f"Date: {self.date_first_shipment.strftime('%B %d, %Y')}")

        doc.save(str(filepath))

    def gen_employee_list(self):
        """Generate Employee_List_Phantom_Logistics.docx with table."""
        filepath = self.output_dir / "Employee_List_Phantom_Logistics.docx"
        doc = Document()

        # Title
        title = doc.add_heading("Phantom Logistics LLC", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_heading("Employee Roster - 2024", 2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Last Updated: {self.date_incorporation.strftime('%B %d, %Y')}")
        doc.add_paragraph("")

        # Employee table
        table = doc.add_table(rows=5, cols=4)
        table.style = 'Light Grid Accent 1'

        # Header row
        headers = ["Name", "Position", "Department", "Contact"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Employee data
        employees = [
            ("Sarah Chen", "Chief Executive Officer", "Executive", "(310) 555-0123"),
            ("David Kim", "Warehouse Manager", "Operations", "(424) 555-0789"),
            ("John Doe", "Chief Financial Officer", "Finance", "(310) 555-0199"),
            ("Maria Lopez", "Administrative Assistant", "Administration", "(310) 555-0145"),
        ]

        for i, (name, position, dept, contact) in enumerate(employees, start=1):
            table.rows[i].cells[0].text = name
            table.rows[i].cells[1].text = position
            table.rows[i].cells[2].text = dept
            table.rows[i].cells[3].text = contact

        doc.add_paragraph("")

        # Company info
        doc.add_heading("Company Information", 2)
        doc.add_paragraph("Registered Address:\n456 Commerce Blvd\nLong Beach, CA 90802")
        doc.add_paragraph("")
        doc.add_paragraph("Offshore Registration:\nCayman Islands (Certificate #CI-2024-001523)")
        doc.add_paragraph("")
        doc.add_paragraph("Primary Contact: Sarah Chen, CEO\nEmail: sarah.chen@phantomlogistics.com")

        doc.save(str(filepath))

    def gen_final_report(self):
        """Generate Final_Report_Customs_Investigation.docx."""
        filepath = self.output_dir / "Final_Report_Customs_Investigation.docx"
        doc = Document()

        # Title page
        title = doc.add_heading("U.S. CUSTOMS AND BORDER PROTECTION", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_heading("INVESTIGATIVE SUMMARY REPORT", 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")
        doc.add_paragraph(f"Case Number: CBP-2024-LA-0523")
        doc.add_paragraph(f"Report Date: {self.date_final_report.strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Classification: CONFIDENTIAL")
        doc.add_paragraph("")
        doc.add_paragraph("_" * 60)
        doc.add_paragraph("")

        # Executive summary
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(
            "This report summarizes the investigation into Phantom Logistics LLC and "
            "associated individuals for suspected smuggling operations, customs fraud, "
            "and public corruption. The investigation, initiated following a whistleblower "
            f"complaint on {self.date_whistleblower.strftime('%B %d, %Y')}, has resulted in arrests and the seizure of "
            "contraband electronics valued at over $2 million."
        )

        doc.add_paragraph("")

        # Key findings
        doc.add_heading("Key Findings", 1)
        doc.add_paragraph("1. SMUGGLING OPERATION: Phantom Logistics LLC operated as a front company for "
                         "smuggling high-value electronics components (advanced microchips) from Hong Kong "
                         "to the United States.")
        doc.add_paragraph("")
        doc.add_paragraph("2. CUSTOMS FRAUD: False declarations on customs forms significantly underreported "
                         "the value of imported goods. Declared value: $250,000. Actual value: $2.5M+.")
        doc.add_paragraph("")
        doc.add_paragraph("3. PUBLIC CORRUPTION: Port Inspector Marcus Rivera (Badge #4521) accepted bribes "
                         "totaling approximately $50,000 to expedite container clearances and falsify "
                         "inspection reports.")
        doc.add_paragraph("")
        doc.add_paragraph("4. FINANCIAL CRIMES: Wire transfers to offshore accounts in the Cayman Islands "
                         "totaling $2.5 million identified as proceeds from smuggling operations.")

        doc.add_paragraph("")

        # Individuals involved
        doc.add_heading("Individuals Involved", 1)

        doc.add_heading("Sarah Chen (Arrested)", 2)
        doc.add_paragraph("Role: CEO of Phantom Logistics LLC")
        doc.add_paragraph("Charges: Smuggling, Customs Fraud, Money Laundering, Conspiracy")
        doc.add_paragraph("Status: Detained pending trial, bail denied")

        doc.add_heading("Marcus Rivera (Suspended)", 2)
        doc.add_paragraph("Role: Port Inspector, Port of Los Angeles")
        doc.add_paragraph("Charges: Bribery, Official Misconduct, Conspiracy")
        doc.add_paragraph("Status: Suspended without pay, under investigation")

        doc.add_heading("David Kim (Cooperating Witness)", 2)
        doc.add_paragraph("Role: Warehouse Manager, Phantom Logistics")
        doc.add_paragraph("Status: Granted immunity in exchange for testimony")
        doc.add_paragraph("Notes: Original whistleblower, provided critical evidence")

        doc.add_paragraph("")

        # Evidence
        doc.add_heading("Evidence Collected", 1)
        doc.add_paragraph("• Container C-999: Seized at Long Beach facility, contained 500 units of "
                         "restricted microchips (export-controlled technology)")
        doc.add_paragraph("• Financial records: Wire transfer receipts, invoices, bank statements")
        doc.add_paragraph("• Email communications: Correspondence between Chen and Rivera "
                         "discussing 'special arrangements'")
        doc.add_paragraph("• Warehouse logs: Documentation of irregular clearance procedures")
        doc.add_paragraph("• Surveillance footage: Video of cash exchange between Chen and Rivera")

        doc.add_paragraph("")

        # Timeline
        doc.add_heading("Timeline of Events", 1)
        timeline = [
            (self.date_incorporation, "Phantom Logistics LLC incorporated in Cayman Islands"),
            (self.date_first_shipment, "Container C-999 arrives at Port of Los Angeles"),
            (self.date_whistleblower, "Anonymous whistleblower complaint filed"),
            (self.date_inspection, "Inspector Rivera files falsified inspection report"),
            (datetime(2024, 4, 14), "Cash payment ($50K) from Chen to Rivera observed"),
            (self.date_wire_transfer, "Wire transfer ($2.5M) to Cayman Islands account"),
            (self.date_second_shipment, "Second container flagged, investigation initiated"),
            (datetime(2024, 5, 10), "Search warrants executed, arrests made"),
        ]

        for date, event in timeline:
            doc.add_paragraph(f"• {date.strftime('%B %d, %Y')}: {event}")

        doc.add_paragraph("")

        # Outcome
        doc.add_heading("Outcome", 1)
        doc.add_paragraph("Assets Seized: $2,500,000 (frozen offshore accounts)")
        doc.add_paragraph("Contraband Seized: 500 microchip units (valued at $2.1M)")
        doc.add_paragraph("Arrests: 2 (Sarah Chen, Marcus Rivera)")
        doc.add_paragraph("Ongoing Investigations: Links to additional shell companies under review")

        doc.add_paragraph("")
        doc.add_paragraph("_" * 60)
        doc.add_paragraph("")
        doc.add_paragraph("Report compiled by: Special Agent Johnson, CBP Internal Affairs")
        doc.add_paragraph(f"Date: {self.date_final_report.strftime('%B %d, %Y')}")
        doc.add_paragraph("Classification: CONFIDENTIAL - Law Enforcement Sensitive")

        doc.save(str(filepath))

    # ========== Email Generators ==========

    def gen_email_thread_c999(self):
        """Generate Email_Thread_Shipment_C999.eml."""
        filepath = self.output_dir / "Email_Thread_Shipment_C999.eml"

        msg = MIMEMultipart()
        msg['From'] = "sarah.chen@phantomlogistics.com"
        msg['To'] = "m.rivera@portofla.gov"
        msg['Subject'] = "Re: Container C-999 - Urgent"
        msg['Date'] = email.utils.formatdate(localtime=True, usegmt=False)
        msg['Message-ID'] = email.utils.make_msgid(domain="phantomlogistics.com")

        body = f"""Hi Marcus,

Thanks for taking care of Container C-999 yesterday. As discussed, this is special cargo that needs to be handled off-books.

The shipment arrived on schedule (Feb 10) and we need it cleared ASAP. Same arrangement as last time - I'll have the cash ready for you at our usual spot.

Please make sure there's no paper trail for this one. Mark it as "routine inspection - no issues found" in your report.

The next shipment (C-1001) arrives March 20th. Let's keep this arrangement confidential.

Best,
Sarah Chen
CEO, Phantom Logistics LLC
(310) 555-0123

---
CONFIDENTIAL: This email and any attachments are intended solely for the use of the individual or entity to whom they are addressed.
"""

        msg.attach(MIMEText(body, 'plain'))

        with open(filepath, 'w') as f:
            f.write(msg.as_string())

    def gen_email_chain_rivera_chen(self):
        """Generate Email_Chain_Rivera_Chen.eml (multi-message thread)."""
        filepath = self.output_dir / "Email_Chain_Rivera_Chen.eml"

        # Create a multi-part message with thread history
        msg = MIMEMultipart()
        msg['From'] = "m.rivera@portofla.gov"
        msg['To'] = "sarah.chen@phantomlogistics.com"
        msg['Subject'] = "Re: Re: Re: Ongoing Arrangements"
        msg['Date'] = email.utils.formatdate(localtime=True, usegmt=False)

        body = f"""Sarah,

Someone's asking questions. One of the warehouse guys has been snooping around. Be careful with the next shipment.

- Marcus

---

On March 10, 2024, Marcus Rivera <m.rivera@portofla.gov> wrote:
> Our arrangement is working well. Next shipment needs same treatment.

---

On February 20, 2024, Sarah Chen <sarah.chen@phantomlogistics.com> wrote:
>> Marcus,
>>
>> Glad to hear everything went smoothly. The payment has been processed as discussed.
>>
>> For the next container (arriving Feb 28), we'll need the same level of... discretion.
>>
>> Best,
>> Sarah

---

On February 5, 2024, Marcus Rivera <m.rivera@portofla.gov> wrote:
>>> Sarah,
>>>
>>> Everything went through without a hitch. Your package cleared customs this morning.
>>>
>>> Looking forward to our continued partnership.
>>>
>>> M.R.
"""

        msg.attach(MIMEText(body, 'plain'))

        with open(filepath, 'w') as f:
            f.write(msg.as_string())

    # ========== Text File Generators ==========

    def gen_warehouse_inventory(self):
        """Generate Warehouse_Inventory_Log.txt (TSV format)."""
        filepath = self.output_dir / "Warehouse_Inventory_Log.txt"

        content = """WAREHOUSE 7B INVENTORY LOG
Long Beach, CA 90802
Logged by: David Kim, Warehouse Manager

Date\t\tTime\tContainer ID\tAction\t\tNotes
================================================================================
2024-02-01\t08:15\tC-888\t\tArrival\t\tStandard processing
2024-02-01\t14:30\tC-888\t\tInspection\tPassed - cleared for release
2024-02-02\t09:00\tC-888\t\tDeparture\tDelivered to consignee
2024-02-10\t11:45\tC-999\t\tArrival\t\tPriority tag - expedited clearance
2024-02-10\t12:00\tC-999\t\tInspection\tInspector Rivera - cleared immediately
2024-02-10\t12:15\tC-999\t\tDeparture\tFast-tracked (no physical inspection?)
2024-02-15\t07:30\tC-1000\t\tArrival\t\tStandard processing
2024-02-15\t10:00\tC-1000\t\tInspection\tPassed - cleared for release
2024-02-16\t08:00\tC-1000\t\tDeparture\tDelivered to consignee
2024-02-28\t13:20\tC-1001\t\tArrival\t\tPriority tag - expedited clearance
2024-02-28\t13:30\tC-1001\t\tInspection\tInspector Rivera - cleared immediately
2024-02-28\t13:45\tC-1001\t\tDeparture\tFast-tracked again (suspicious)
2024-03-05\t09:15\tC-1002\t\tArrival\t\tStandard processing
2024-03-05\t14:00\tC-1002\t\tInspection\tPassed - cleared for release
2024-03-06\t08:30\tC-1002\t\tDeparture\tDelivered to consignee
2024-03-20\t10:00\tC-1003\t\tArrival\t\tPriority tag - UNUSUAL ACTIVITY
2024-03-20\t10:10\tC-1003\t\tInspection\tInspector Rivera - no log entry?
2024-03-20\t10:20\tC-1003\t\tDeparture\tCleared without proper paperwork
2024-04-01\t15:45\tC-999\t\tRe-inspection\tRivera filed report - claims no issues
2024-04-15\t08:00\tC-1004\t\tArrival\t\tStandard processing
2024-04-15\t11:30\tC-1004\t\tInspection\tPassed - cleared for release
2024-04-16\t07:00\tC-1004\t\tDeparture\tDelivered to consignee

================================================================================
NOTES:
- All containers with ID pattern C-999 through C-1003 show irregular clearance
- Inspector Rivera bypassing standard inspection protocols
- Need to report this to supervisor
- Containers cleared in <30 minutes (normal time: 2-4 hours)

Logged by: David Kim
Position: Warehouse Manager
Contact: (424) 555-0789
"""

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

    def gen_phone_records(self):
        """Generate Phone_Records_Extract.txt (CSV-like CDR)."""
        filepath = self.output_dir / "Phone_Records_Extract.txt"

        content = """CALL DETAIL RECORDS (CDR) EXTRACT
Port Authority Investigation - Case #2024-LA-0523
Date Range: January 1, 2024 - May 1, 2024

Format: Date,Time,From,To,Duration(sec),Type
================================================================================
2024-01-20,14:30:00,(310) 555-0123,(213) 555-0456,180,Outgoing
2024-01-25,09:15:00,(213) 555-0456,(310) 555-0123,240,Outgoing
2024-02-05,11:00:00,(310) 555-0123,(213) 555-0456,300,Outgoing
2024-02-08,16:45:00,(310) 555-0123,(213) 555-0456,420,Outgoing
2024-02-10,12:30:00,(213) 555-0456,(310) 555-0123,180,Outgoing
2024-02-15,10:00:00,(424) 555-0789,(310) 555-0123,120,Outgoing
2024-02-20,13:20:00,(310) 555-0123,(213) 555-0456,360,Outgoing
2024-02-28,14:00:00,(213) 555-0456,(310) 555-0123,240,Outgoing
2024-03-05,09:30:00,(310) 555-0123,(213) 555-0456,180,Outgoing
2024-03-10,15:15:00,(310) 555-0123,(213) 555-0456,480,Outgoing
2024-03-12,08:00:00,(424) 555-0789,(213) 555-0456,90,Outgoing
2024-03-15,17:30:00,(424) 555-0789,(310) 555-0123,300,Outgoing
2024-03-20,10:45:00,(213) 555-0456,(310) 555-0123,120,Outgoing
2024-04-01,16:00:00,(310) 555-0123,(213) 555-0456,240,Outgoing
2024-04-14,19:30:00,(310) 555-0123,(213) 555-0456,600,Outgoing
2024-04-15,10:15:00,(213) 555-0456,(310) 555-0123,180,Outgoing
2024-04-25,14:45:00,(424) 555-0789,(310) 555-0123,420,Outgoing
2024-05-01,11:00:00,(213) 555-0456,(310) 555-0123,300,Outgoing

================================================================================
SUBSCRIBER INFORMATION:
(310) 555-0123 - Sarah Chen (Phantom Logistics LLC)
(213) 555-0456 - Marcus Rivera (Port Authority)
(424) 555-0789 - David Kim (Warehouse Manager)

ANALYSIS NOTES:
- High frequency of calls between Chen and Rivera (suspicious)
- Spike in call duration around container arrival dates
- Kim called Rivera on March 12 (possible confrontation?)
- Pattern suggests coordination of illegal activities

Extracted by: Special Agent Johnson, CBP
Date: May 2, 2024
Classification: CONFIDENTIAL
"""

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

    def gen_gps_coordinates(self):
        """Generate GPS_Coordinates_Warehouse.txt."""
        filepath = self.output_dir / "GPS_Coordinates_Warehouse.txt"

        content = """GPS LOCATION DATA
Project Phantom Investigation - Key Locations
================================================================================

WAREHOUSE 7B - PRIMARY FACILITY
Coordinates: 33.7701° N, 118.1937° W
Address: 456 Commerce Blvd, Long Beach, CA 90802
Description: Phantom Logistics LLC warehouse facility
Notes: Container storage and processing location

PIER 9 - PORT OF LOS ANGELES
Coordinates: 33.7434° N, 118.2719° W
Address: Terminal C, Pier 9, San Pedro, CA 90731
Description: Container offloading terminal
Notes: Inspector Rivera's primary patrol area

MEETING LOCATION - CASH EXCHANGE SITE
Coordinates: 33.7550° N, 118.2000° W
Address: Parking lot near Port Administration Building
Description: Secluded area identified in surveillance
Notes: Site of April 14 cash payment ($50K) from Chen to Rivera

ACME IMPORTS INC. - FRONT COMPANY
Coordinates: 34.0522° N, 118.2437° W
Address: 789 Trade Center Drive, Los Angeles, CA 90012
Description: Registered address for Acme Imports
Notes: Suspected money laundering front

CAYMAN ISLANDS REGISTRATION
Coordinates: 19.2866° N, 81.3744° W
Address: 123 Offshore Drive, George Town, Grand Cayman
Description: Phantom Logistics LLC incorporation address
Notes: Offshore jurisdiction, tax haven

================================================================================
GEOSPATIAL ANALYSIS:

Distance Matrix:
- Warehouse 7B to Pier 9: 3.2 miles (5.1 km)
- Warehouse 7B to Meeting Location: 1.8 miles (2.9 km)
- Pier 9 to Meeting Location: 1.5 miles (2.4 km)

Pattern: All key locations within 3-mile radius, suggesting coordinated operation

Compiled by: CBP Geospatial Intelligence Unit
Date: May 2, 2024
"""

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

    # ========== Image Generators ==========

    def gen_wire_transfer_image(self):
        """Generate Bank_Wire_Transfer_Receipt.jpg (photo of receipt)."""
        filepath = self.output_dir / "Bank_Wire_Transfer_Receipt.jpg"

        # Create image (simulated receipt photo)
        img = Image.new('RGB', (800, 1000), color='white')
        draw = ImageDraw.Draw(img)

        # Try to use a standard font, fallback to default
        try:
            font_large = ImageFont.truetype("arial.ttf", 28)
            font_medium = ImageFont.truetype("arial.ttf", 20)
            font_small = ImageFont.truetype("arial.ttf", 16)
        except:
            font_large = ImageFont.load_default()
            font_medium = ImageFont.load_default()
            font_small = ImageFont.load_default()

        # Header
        y = 50
        draw.text((400, y), "FIRST NATIONAL BANK", fill='black', anchor='mm', font=font_large)
        y += 40
        draw.text((400, y), "Wire Transfer Confirmation", fill='black', anchor='mm', font=font_medium)

        y += 80

        # Receipt details
        lines = [
            ("Transaction ID:", "WTR-2024-0415-8877"),
            ("", ""),
            ("Date:", self.date_wire_transfer.strftime('%B %d, %Y')),
            ("Time:", "10:23:15 AM PST"),
            ("", ""),
            ("FROM:", ""),
            ("Account Name:", "Acme Imports Inc."),
            ("Account Number:", "**** **** **** 7890"),
            ("Bank:", "First National Bank of California"),
            ("Address:", "Los Angeles, CA 90012"),
            ("", ""),
            ("TO:", ""),
            ("Beneficiary:", "Phantom Offshore Holdings"),
            ("Account Number:", "CAY-99-8877-XX"),
            ("Bank:", "Cayman International Bank"),
            ("Address:", "George Town, Grand Cayman"),
            ("", ""),
            ("AMOUNT:", "$2,500,000.00 USD"),
            ("", ""),
            ("Purpose:", "Consulting Services - Invoice ACM-2024-0415"),
            ("", ""),
            ("Status:", "COMPLETED"),
            ("", ""),
            ("Reference Number:", "FNB-WTR-415-2024-8877"),
            ("", ""),
            ("", ""),
            ("This is an official bank record.", ""),
            ("For inquiries, call: 1-800-555-BANK", ""),
        ]

        for label, value in lines:
            if label == "AMOUNT:":
                # Highlight amount
                draw.rectangle([(50, y-5), (750, y+25)], outline='red', width=2)
            if label and value:
                draw.text((80, y), label, fill='black', font=font_small)
                draw.text((350, y), value, fill='black', font=font_small)
            elif label:
                draw.text((80, y), label, fill='black', font=font_medium)
            y += 30

        # Save with slight JPEG compression (simulate photo)
        img.save(str(filepath), 'JPEG', quality=85)

    def gen_handwritten_note(self):
        """Generate Handwritten_Note_Meeting.jpg (simulated handwriting)."""
        filepath = self.output_dir / "Handwritten_Note_Meeting.jpg"

        # Create image (simulated handwritten note)
        img = Image.new('RGB', (600, 400), color='#FFF8DC')  # Cream paper color
        draw = ImageDraw.Draw(img)

        # Try to use a script/cursive-style font, fallback to default
        try:
            # Try common script fonts
            for font_name in ['Brush Script MT', 'Comic Sans MS', 'Lucida Handwriting']:
                try:
                    font = ImageFont.truetype(font_name, 32)
                    break
                except:
                    continue
            else:
                font = ImageFont.load_default()
        except:
            font = ImageFont.load_default()

        # Handwritten text (slightly rotated for realism would be ideal, but keeping simple)
        y = 80
        lines = [
            "Meeting with M.R.",
            "$50K cash",
            "April 14 - Pier 9",
            "Usual spot",
            "- S.C."
        ]

        for line in lines:
            # Add slight random offset to simulate handwriting irregularity
            x_offset = 50 + (hash(line) % 20)
            draw.text((x_offset, y), line, fill='#000080', font=font)
            y += 50

        # Add some "paper texture" (subtle noise)
        # Save
        img.save(str(filepath), 'JPEG', quality=90)

    # ========== README Generator ==========

    def gen_readme(self):
        """Generate README.md for the tutorial case."""
        filepath = self.output_dir / "README.md"

        content = f"""# Project Phantom - Tutorial Case

This is a **fictional investigative case** designed to demonstrate all features of ArkhamMirror.

## 📋 Case Summary

**Phantom Logistics LLC** is a shell company suspected of maritime smuggling operations. The case involves:

* **Smuggling**: Contraband electronics concealed in shipping containers
* **Corruption**: Port inspector accepting bribes to bypass inspections
* **Fraud**: False customs declarations and offshore money laundering
* **Whistleblower**: Warehouse manager exposes the operation

## 📁 Files Included (15 Documents)

| # | Filename | Type | Purpose |
|---|----------|------|---------|
| 1 | Company_Formation_Certificate.pdf | PDF (scanned) | OCR test, entity extraction |
| 2 | Email_Thread_Shipment_C999.eml | Email | Anomaly detection ("off-books") |
| 3 | Shipping_Manifest_C999.docx | DOCX | Table extraction, entity linking |
| 4 | Warehouse_Inventory_Log.txt | Text (TSV) | Regex search (container IDs) |
| 5 | Whistleblower_Complaint.pdf | PDF (typed) | Anomaly detection ("bribe") |
| 6 | Port_Inspection_Report_April1.pdf | PDF (form) | Timeline extraction |
| 7 | Bank_Wire_Transfer_Receipt.jpg | Image | OCR (photo), regex (account #) |
| 8 | Handwritten_Note_Meeting.jpg | Image | OCR (handwriting) via Qwen-VL |
| 9 | Email_Chain_Rivera_Chen.eml | Email | Email parsing, entity graph |
| 10 | Customs_Declaration_Form_C999.pdf | PDF | Timeline, entity extraction |
| 11 | Employee_List_Phantom_Logistics.docx | DOCX (table) | Table extraction, entities |
| 12 | Phone_Records_Extract.txt | Text (CSV) | Regex search (phone numbers) |
| 13 | GPS_Coordinates_Warehouse.txt | Text | Geolocation mapping |
| 14 | Invoice_Acme_to_Phantom.pdf | PDF | Financial analysis |
| 15 | Final_Report_Customs_Investigation.docx | DOCX | Search, summary generation |

## 🚀 Quick Start

### 1. Start ArkhamMirror

```bash
# Terminal 1: Start infrastructure
cd arkham_mirror
docker compose up -d

# Terminal 2: Start RQ worker
cd arkham_mirror
.\\venv\\Scripts\\activate  # Windows
python run_rq_worker.py

# Terminal 3: Start Reflex app
cd arkham_reflex
reflex run
```

### 2. Ingest Documents

1. Open browser: `http://localhost:3000`
2. Navigate to **"Ingest"** page (sidebar)
3. Drag and drop all 15 files from this directory
4. Click **"Start Processing"**
5. Wait ~10-15 minutes for processing to complete

### 3. Explore Features

| Feature | What to Try |
|---------|-------------|
| **Search** (`/`) | Search "Who received payments?" or "Container C-999" |
| **Entity Dedup** (`/entity-dedup`) | Find "Sarah Chen" vs. "S. Chen" duplicates |
| **Timeline** (`/timeline`) | Spot March 16-30 suspicious gap |
| **Graph** (`/graph`) | See Sarah Chen ↔ Marcus Rivera connection |
| **Map** (`/map`) | View 3 location markers (Warehouse, Pier, Meeting) |
| **Anomalies** (`/anomalies`) | See "off-books" and "bribe" flagged |
| **Regex** (`/regex-search`) | Find phone numbers, container IDs (C-999) |

## 🎓 Tutorial Walkthrough

For detailed step-by-step instructions, see:

**TUTORIAL_DATASET_DESIGN.md** in the project root

### Key Search Queries

* **"Who received bribes?"** → Marcus Rivera
* **"Container C-999"** → Shipping manifests, emails
* **"wire transfer Cayman"** → Offshore accounts
* **"whistleblower"** → Anonymous complaint
* **"Sarah Chen"** → CEO of Phantom Logistics

## 📊 Expected Results

After processing all 15 documents:

* **Total Documents**: 15
* **Entities Extracted**: ~17 (8 PERSON, 4 ORG, 5 GPE)
* **Canonical Entities**: ~14 (after deduplication)
* **Timeline Events**: 10+
* **Temporal Gaps**: 1 suspicious (March 16-30)
* **Anomalies Detected**: 3-5
* **Geolocation Markers**: 3
* **Extracted Tables**: 1 (employee roster)

### Regex Patterns to Try

* **Phone Numbers**: `\\(\\d{{3}}\\) \\d{{3}}-\\d{{4}}`
  * Finds: (310) 555-0123, (213) 555-0456, (424) 555-0789
* **Container IDs**: `C-\\d{{3}}`
  * Finds: C-999, C-1000, C-1001, etc.
* **Offshore Account**: `CAY-\\d{{2}}-\\d{{4}}-[A-Z]{{2}}`
  * Finds: CAY-99-8877-XX

## 🔍 What This Case Demonstrates

### ✅ Core Features

* **Multi-format ingestion**: PDF, DOCX, EML, TXT, JPG
* **Hybrid OCR**: PaddleOCR (fast) + Qwen-VL (handwriting)
* **Semantic search**: Meaning-based, not keyword matching
* **Entity extraction**: Automatic NER (people, orgs, places)
* **Entity deduplication**: Fuzzy matching and merging

### ✅ Advanced Features

* **Timeline analysis**: Chronological event extraction
* **Temporal gap detection**: Suspicious date ranges
* **Entity relationship graphs**: Visualize connections
* **Geospatial mapping**: Plot locations on interactive map
* **Anomaly detection**: Suspicious language patterns
* **Regex pattern search**: Structured data extraction
* **Table extraction**: Parse structured data from docs

## 💡 Tips

* **Entity Deduplication**: Manually merge "Sarah Chen", "S. Chen", "Dr. Chen"
* **Timeline Gap**: Look for March 16-30 (no docs between whistleblower complaint and inspection)
* **Graph View**: Filter by entity type to focus on PERSON connections
* **Anomalies**: Check why emails are flagged (keywords like "off-books", "confidential")
* **Export**: Try CSV export on search results, timeline, graph, and anomalies

## ⚠️ Disclaimer

This is **100% fictional data**. All names, companies, locations, and events are invented for tutorial purposes. Any resemblance to real persons or entities is coincidental.

## 📚 Learn More

* **User Guide**: `arkham_mirror/docs/USER_GUIDE.md`
* **Full Tutorial**: `TUTORIAL_DATASET_DESIGN.md`
* **Test Plan**: `E2E_TEST_PLAN.md`

---

**Generated**: {datetime.now().strftime('%B %d, %Y')}
**ArkhamMirror Version**: v1.0+
"""

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)


def main():
    """Main entry point."""
    generator = TutorialDatasetGenerator()
    generator.generate_all()


if __name__ == "__main__":
    main()
