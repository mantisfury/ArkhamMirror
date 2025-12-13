#!/usr/bin/env python3
"""
Generate Extended Tutorial Dataset for ArkhamMirror

Extends the basic tutorial with additional documents to showcase ALL features,
especially the advanced AI-powered investigation tools added in Phase 4:
- Contradiction Detection
- Red Flag Discovery
- Hidden Content Detection
- Big Picture Engine
- Narrative Reconstruction
- Speculation Mode
- Fact Comparison
- Fingerprint Detection
- Timeline Merging

Usage:
    python scripts/generate_extended_tutorial.py

This script adds 10 additional documents to the basic tutorial dataset.
"""

import os
import sys
from pathlib import Path
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import email.utils
import random

# Add arkham_mirror to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / "arkham_mirror"))

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image, ImageDraw, ImageFont
except ImportError as e:
    print(f"ERROR: Missing required library: {e}")
    print("\nInstall dependencies:")
    print("pip install reportlab python-docx pillow")
    sys.exit(1)


class ExtendedTutorialGenerator:
    """Generates additional documents to test ALL ArkhamMirror features."""

    def __init__(self, output_dir: str = "data/tutorial_case"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Base dates
        self.date_incorporation = datetime(2024, 1, 15)
        self.date_first_shipment = datetime(2024, 2, 10)
        self.date_whistleblower = datetime(2024, 3, 15)
        self.date_inspection = datetime(2024, 4, 1)
        self.date_wire_transfer = datetime(2024, 4, 15)
        self.date_final_report = datetime(2024, 5, 15)

        print(f"[*] Output directory: {self.output_dir.absolute()}")

    def generate_all(self):
        """Generate all extended tutorial documents."""
        print("\n[*] Generating Extended Tutorial Documents...\n")

        generators = [
            # Documents for CONTRADICTION DETECTION
            (1, "Deposition_Sarah_Chen.pdf", self.gen_chen_deposition),
            (2, "Deposition_Marcus_Rivera.pdf", self.gen_rivera_deposition),
            # Documents for RED FLAG DISCOVERY
            (3, "Financial_Transactions_Q1_2024.txt", self.gen_financial_transactions),
            (4, "Shell_Company_Registry.pdf", self.gen_shell_company_registry),
            # Documents for HIDDEN CONTENT / METADATA FORENSICS
            (5, "Backdated_Contract.pdf", self.gen_backdated_contract),
            (6, "Suspicious_PDF_Analysis.txt", self.gen_metadata_notes),
            # Documents for BIG PICTURE / NARRATIVE
            (7, "Investigation_Timeline_Summary.docx", self.gen_investigation_timeline),
            (8, "Key_Players_Profile.docx", self.gen_key_players),
            # Documents for FINGERPRINT DETECTION (near-duplicates)
            (9, "Invoice_Template_v1.pdf", self.gen_invoice_template_v1),
            (10, "Invoice_Template_v2.pdf", self.gen_invoice_template_v2),
            # Documents for SPECULATION MODE / FACT COMPARISON
            (11, "Contradicting_Witness_Statement.pdf", self.gen_witness_statement),
            (12, "Missing_Period_Analysis.txt", self.gen_missing_period),
        ]

        for num, filename, generator_func in generators:
            try:
                print(f"[{num}/12] Generating {filename}...")
                generator_func()
                print(f"    [OK] Created: {filename}")
            except Exception as e:
                print(f"    [ERROR] {e}")
                import traceback

                traceback.print_exc()

        # Update README
        self.gen_extended_readme()
        print(f"\n[OK] Extended documents generated!")
        print(f"[*] Total files: 27 (15 basic + 12 extended)")

    # ========== CONTRADICTION DETECTION Documents ==========

    def gen_chen_deposition(self):
        """Generate Chen's deposition with specific claims."""
        filepath = self.output_dir / "Deposition_Sarah_Chen.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width / 2, height - inch, "SWORN DEPOSITION")
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(
            width / 2,
            height - 1.3 * inch,
            "IN THE MATTER OF: U.S. v. PHANTOM LOGISTICS LLC",
        )

        c.setFont("Helvetica", 10)
        c.drawString(inch, height - 1.8 * inch, f"Deponent: Sarah Chen")
        c.drawString(
            inch,
            height - 2 * inch,
            f"Date: {self.date_final_report.strftime('%B %d, %Y')}",
        )

        # Key claims (some will CONTRADICT Rivera's deposition)
        content = """
EXAMINATION BY MR. PROSECUTOR:

Q: Ms. Chen, how many times did you meet with Inspector Rivera?
A: I only met him once, at a logistics conference in January 2024.

Q: Did you ever provide cash payments to Inspector Rivera?
A: Absolutely not. I have never paid anyone in cash for services.

Q: The records show a wire transfer of $2.5 million to the Cayman Islands. Can you explain?
A: That was a legitimate consulting fee for market analysis services provided by Acme Imports.

Q: Who authorized Container C-999 to bypass standard inspection?
A: I have no knowledge of any containers bypassing inspection. All our shipments went through normal customs procedures.

Q: When did you first learn about the investigation?
A: I was completely blindsided when agents arrived at my office on May 10th.

Q: Did you ever communicate with Inspector Rivera via email about "special arrangements"?
A: Those emails are taken out of context. We were discussing scheduling arrangements for container pickups.

Q: The whistleblower complaint mentions cash exchanges. Your response?
A: I don't know who filed that complaint, but those allegations are completely false.

[DEPOSITION CONTINUED ON NEXT PAGE]
"""

        y = height - 2.5 * inch
        c.setFont("Helvetica", 10)
        for line in content.strip().split("\n"):
            if y < 1.5 * inch:
                c.showPage()
                y = height - inch
                c.setFont("Helvetica", 10)
            c.drawString(inch, y, line)
            y -= 0.2 * inch

        c.save()

    def gen_rivera_deposition(self):
        """Generate Rivera's deposition with CONTRADICTING claims."""
        filepath = self.output_dir / "Deposition_Marcus_Rivera.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width / 2, height - inch, "SWORN DEPOSITION")
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(
            width / 2,
            height - 1.3 * inch,
            "IN THE MATTER OF: U.S. v. PHANTOM LOGISTICS LLC",
        )

        c.setFont("Helvetica", 10)
        c.drawString(inch, height - 1.8 * inch, f"Deponent: Marcus Rivera")
        c.drawString(
            inch,
            height - 2 * inch,
            f"Date: {self.date_final_report.strftime('%B %d, %Y')}",
        )

        # CONTRADICTING statements
        content = """
EXAMINATION BY MR. PROSECUTOR:

Q: Mr. Rivera, how often did you meet with Sarah Chen?
A: We met multiple times - at least five or six occasions between January and April.

Q: Did you receive any payments from Ms. Chen or Phantom Logistics?
A: [After consulting with counsel] Yes, I received cash payments totaling approximately $50,000 over several months.

Q: Who authorized Container C-999 to bypass inspection?
A: I did. Ms. Chen specifically asked me to expedite clearance and provided payment for this service.

Q: Were these payments documented anywhere?
A: No. Ms. Chen was very clear that everything should be "off the books" - her exact words.

Q: When did your arrangement with Phantom Logistics begin?
A: Sarah Chen approached me at a conference in January. We had our first... transaction... 
   shortly after that, around late January.

Q: How were the cash exchanges conducted?
A: Usually at Pier 9, in the parking lot. Sometimes at a coffee shop on Ocean Boulevard.

Q: Did Ms. Chen ever express concern about being caught?
A: Yes, after the whistleblower complaint was filed, she became very nervous. She said 
   we needed to "destroy any paper trail" and that she had "connections" to handle it.

[DEPOSITION CONTINUED ON NEXT PAGE]
"""

        y = height - 2.5 * inch
        c.setFont("Helvetica", 10)
        for line in content.strip().split("\n"):
            if y < 1.5 * inch:
                c.showPage()
                y = height - inch
                c.setFont("Helvetica", 10)
            c.drawString(inch, y, line)
            y -= 0.2 * inch

        c.save()

    # ========== RED FLAG DISCOVERY Documents ==========

    def gen_financial_transactions(self):
        """Generate financial records with RED FLAG patterns."""
        filepath = self.output_dir / "Financial_Transactions_Q1_2024.txt"

        # Include RED FLAG patterns:
        # - Round numbers ($9,000 - structuring)
        # - Multiple transactions just below $10K (smurfing)
        # - Suspiciously regular amounts

        content = """PHANTOM LOGISTICS LLC - FINANCIAL TRANSACTION LEDGER
Account: First National Bank #123456789
Period: January 1 - March 31, 2024

Date        | Transaction ID | Type      | Amount     | Description              | Recipient/Sender
------------+----------------+-----------+------------+--------------------------+------------------
2024-01-05  | TXN-001        | WIRE OUT  | $9,500.00  | Consulting Fee           | Cayman Trust Co.
2024-01-08  | TXN-002        | WIRE OUT  | $9,800.00  | Advisory Services        | Offshore Holdings Ltd.
2024-01-12  | TXN-003        | WIRE OUT  | $9,200.00  | Strategic Planning       | Cayman Trust Co.
2024-01-15  | TXN-004        | WIRE OUT  | $9,999.00  | Market Research          | Pacific Advisory
2024-01-18  | TXN-005        | DEPOSIT   | $50,000.00 | Client Payment           | Global Maritime
2024-01-22  | TXN-006        | WIRE OUT  | $9,700.00  | Consulting Fee           | Offshore Holdings Ltd.
2024-01-25  | TXN-007        | WIRE OUT  | $9,900.00  | Advisory Services        | Cayman Trust Co.
2024-01-28  | TXN-008        | WIRE OUT  | $9,100.00  | Strategic Planning       | Pacific Advisory
2024-02-01  | TXN-009        | DEPOSIT   | $100,000.00| Retainer Fee             | Acme Imports Inc.
2024-02-05  | TXN-010        | WIRE OUT  | $9,600.00  | Consulting Fee           | Cayman Trust Co.
2024-02-08  | TXN-011        | WIRE OUT  | $9,400.00  | Advisory Services        | Offshore Holdings Ltd.
2024-02-12  | TXN-012        | WIRE OUT  | $9,850.00  | Market Research          | Pacific Advisory
2024-02-15  | TXN-013        | WIRE OUT  | $9,750.00  | Strategic Planning       | Cayman Trust Co.
2024-02-18  | TXN-014        | DEPOSIT   | $250,000.00| Container Fees           | Global Maritime
2024-02-22  | TXN-015        | WIRE OUT  | $9,950.00  | Consulting Fee           | Offshore Holdings Ltd.
2024-02-25  | TXN-016        | WIRE OUT  | $9,550.00  | Advisory Services        | Pacific Advisory
2024-03-01  | TXN-017        | DEPOSIT   | $500,000.00| Project Alpha Payment    | Undisclosed Client
2024-03-05  | TXN-018        | WIRE OUT  | $9,300.00  | Consulting Fee           | Cayman Trust Co.
2024-03-08  | TXN-019        | WIRE OUT  | $9,450.00  | Strategic Advisory       | Offshore Holdings Ltd.
2024-03-12  | TXN-020        | WIRE OUT  | $9,650.00  | Market Analysis          | Pacific Advisory
2024-03-15  | TXN-021        | WIRE OUT  | $9,875.00  | Risk Assessment          | Cayman Trust Co.
2024-03-20  | TXN-022        | WIRE OUT  | $9,125.00  | Consulting Fee           | Offshore Holdings Ltd.
2024-03-25  | TXN-023        | WIRE OUT  | $9,990.00  | Strategic Planning       | Pacific Advisory
2024-03-28  | TXN-024        | WIRE OUT  | $9,425.00  | Advisory Services        | Cayman Trust Co.

---
SUMMARY:
Total Deposits: $900,000.00
Total Withdrawals (Wire): $175,820.00 (18 transactions averaging $9,767.78)
Remaining Balance: $724,180.00

NOTE: All outgoing wires directed to offshore jurisdictions. 
Multiple transactions below $10,000 reporting threshold detected.
"""

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

    def gen_shell_company_registry(self):
        """Generate shell company information for red flag detection."""
        filepath = self.output_dir / "Shell_Company_Registry.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(
            width / 2, height - inch, "CORPORATE REGISTRY SEARCH RESULTS"
        )
        c.setFont("Helvetica", 10)
        c.drawCentredString(
            width / 2,
            height - 1.3 * inch,
            "Compiled by: FinDiscovery Investigations LLC",
        )

        y = height - 2 * inch
        c.setFont("Helvetica-Bold", 11)
        c.drawString(inch, y, "Subject: Phantom Logistics LLC and Associated Entities")
        y -= 0.4 * inch

        content = [
            "FINDINGS:",
            "",
            "1. PHANTOM LOGISTICS LLC (Cayman Islands)",
            "   - Incorporated: January 15, 2024",
            "   - Type: Exempt Limited Partnership",
            "   - Registered Agent: Cayman Corporate Services",
            "   - Directors: Sarah Chen (sole)",
            "   - No publicly reported revenue",
            "   - RED FLAG: No verifiable business operations in jurisdiction",
            "",
            "2. PHANTOM OFFSHORE HOLDINGS (Cayman Islands)",
            "   - Incorporated: January 20, 2024 (5 days after Phantom Logistics)",
            "   - Type: Holding Company",
            "   - Same registered agent as Phantom Logistics",
            "   - Directors: Unnamed nominee directors",
            "   - RED FLAG: Bearer shares, hidden ownership",
            "",
            "3. ACME IMPORTS INC. (Delaware, USA)",
            "   - Incorporated: December 10, 2023",
            "   - CEO: Listed as 'John Doe' (likely fictitious)",
            "   - Address: Virtual office in Wilmington",
            "   - RED FLAG: No physical operations, minimal filings",
            "",
            "4. CAYMAN TRUST CO. (Cayman Islands)",
            "   - Recipient of 8 wire transfers from Phantom Logistics",
            "   - Total received: $77,675.00",
            "   - NO CORPORATE RECORDS FOUND",
            "   - RED FLAG: Possibly unregistered or dissolved",
            "",
            "5. OFFSHORE HOLDINGS LTD. (British Virgin Islands)",
            "   - Registered: 2019",
            "   - 6 wire transfers from Phantom Logistics: $57,250.00",
            "   - Linked to 47 other shell companies",
            "   - RED FLAG: Part of known layering network",
            "",
            "CONCLUSION:",
            "Strong indicators of corporate layering scheme designed to obscure",
            "the movement of funds and beneficial ownership. Recommend full",
            "forensic accounting investigation.",
        ]

        c.setFont("Helvetica", 10)
        for line in content:
            if y < 1.5 * inch:
                c.showPage()
                y = height - inch
                c.setFont("Helvetica", 10)
            c.drawString(inch, y, line)
            y -= 0.2 * inch

        c.save()

    # ========== HIDDEN CONTENT / METADATA Documents ==========

    def gen_backdated_contract(self):
        """Generate a PDF with suspicious metadata (backdating)."""
        filepath = self.output_dir / "Backdated_Contract.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        # Note: The PDF metadata itself won't be truly manipulated,
        # but the CONTENT describes a backdating scenario

        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width / 2, height - inch, "CONSULTING AGREEMENT")

        c.setFont("Helvetica", 10)
        y = height - 1.5 * inch

        content = [
            "Agreement Date: December 15, 2023",  # Backdated
            "Document Created: May 8, 2024",  # Actual creation (after investigation started)
            "",
            "PARTIES:",
            "1. Phantom Logistics LLC ('Client')",
            "2. Acme Imports Inc. ('Consultant')",
            "",
            "TERMS:",
            "",
            "This agreement establishes the consulting relationship whereby Consultant",
            "shall provide 'strategic advisory services' to Client.",
            "",
            "1. SCOPE OF WORK",
            "   Consultant agrees to provide:",
            "   - Market analysis (unspecified)",
            "   - Strategic planning (unspecified)",
            "   - Risk management (unspecified)",
            "",
            "2. COMPENSATION",
            "   Client shall pay Consultant $2,500,000.00 USD upon execution.",
            "   Payment shall be wired to: Phantom Offshore Holdings, Cayman Islands",
            "",
            "3. TERM",
            "   This agreement is effective as of December 15, 2023 and continues",
            "   until completion of services (undefined).",
            "",
            "[NOTE: This contract was produced AFTER federal investigation began.",
            " Metadata analysis shows creation date of May 2024, despite claiming",
            " execution in December 2023. Document appears to have been created",
            " to justify the $2.5M wire transfer retroactively.]",
            "",
            "SIGNATURES:",
            "",
            "___________________          ___________________",
            "Sarah Chen                   John Doe",
            "CEO, Phantom Logistics       CEO, Acme Imports",
        ]

        for line in content:
            c.drawString(inch, y, line)
            y -= 0.2 * inch

        c.save()

    def gen_metadata_notes(self):
        """Generate metadata forensics analysis notes."""
        filepath = self.output_dir / "Suspicious_PDF_Analysis.txt"

        content = """METADATA FORENSICS ANALYSIS REPORT
Prepared by: Digital Forensics Unit
Case: CBP-2024-LA-0523 (Phantom Logistics)
Date: May 12, 2024

DOCUMENT: Backdated_Contract.pdf

FILE METADATA:
  - Filename: Backdated_Contract.pdf
  - File Size: 45,892 bytes
  - PDF Version: 1.5
  - Producer: Microsoft Print to PDF
  
CREATION METADATA:
  - Creation Date (in metadata): 2023-12-15T14:30:00Z
  - Modification Date: 2024-05-08T09:15:22Z
  
ANOMALIES DETECTED:

1. CREATION/MODIFICATION MISMATCH (CRITICAL)
   The modification date (May 8, 2024) is AFTER the creation date (Dec 15, 2023).
   However, detailed analysis shows the PDF was actually CREATED on May 8, 2024.
   The December 2023 date was manually inserted into metadata.

2. PRODUCER SOFTWARE MISMATCH
   Document claims to be from December 2023, but the "Microsoft Print to PDF"
   version used did not receive that update until February 2024.
   CONCLUSION: Document could not have been created in December 2023.

3. FONT EMBEDDING TIMESTAMPS
   Embedded fonts show registration dates of March 2024.
   CONCLUSION: Fonts updated after purported creation date.

4. DIGITAL SIGNATURE ABSENCE
   No digital signature present, despite company policy (per internal docs)
   requiring signature on contracts > $100,000.

5. TIMEZONE INCONSISTENCY
   Creation metadata shows UTC timezone, but Phantom Logistics is based in
   California (PST/PDT). All other verified documents show local timezone.

FORENSIC CONCLUSION:
This document was created on or around May 8, 2024 - AFTER the federal 
investigation was publicly known - and was retroactively backdated to 
December 2023. This constitutes potential evidence tampering.

RECOMMENDATION:
Add to obstruction of justice charges. Interview document author (File Properties
shows "Author: SChen-PC") regarding document creation timeline.
"""

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

    # ========== BIG PICTURE / NARRATIVE Documents ==========

    def gen_investigation_timeline(self):
        """Generate comprehensive timeline for Big Picture analysis."""
        filepath = self.output_dir / "Investigation_Timeline_Summary.docx"
        doc = Document()

        doc.add_heading("PHANTOM LOGISTICS INVESTIGATION TIMELINE", 0)
        doc.add_heading("Comprehensive Event Chronology", 2)

        events = [
            ("December 10, 2023", "Acme Imports Inc. incorporated in Delaware"),
            (
                "January 15, 2024",
                "Phantom Logistics LLC incorporated in Cayman Islands",
            ),
            (
                "January 20, 2024",
                "Phantom Offshore Holdings registered (same jurisdiction)",
            ),
            (
                "Late January 2024",
                "Sarah Chen and Marcus Rivera allegedly meet at conference",
            ),
            (
                "January - March 2024",
                "Multiple wire transfers ($9K-$10K range) to offshore accounts",
            ),
            ("February 10, 2024", "Container C-999 arrives at Port of Los Angeles"),
            (
                "February 10, 2024",
                "C-999 cleared by Inspector Rivera without physical inspection",
            ),
            (
                "February 28, 2024",
                "Container C-1001 arrives, same expedited processing",
            ),
            (
                "March 10, 2024",
                "Email evidence shows Rivera warning Chen about 'someone asking questions'",
            ),
            (
                "March 15, 2024",
                "Anonymous whistleblower complaint filed with Port Authority",
            ),
            ("March 16-30, 2024", "[SUSPICIOUS GAP - NO DOCUMENTS]"),
            (
                "April 1, 2024",
                "Rivera files official inspection report claiming 'no irregularities'",
            ),
            (
                "April 14, 2024",
                "Alleged $50K cash payment from Chen to Rivera (per handwritten note)",
            ),
            (
                "April 15, 2024",
                "Wire transfer of $2.5M from Acme to Phantom Offshore Holdings",
            ),
            ("May 1, 2024", "Second container flagged by independent customs review"),
            (
                "May 8, 2024",
                "Backdated contract created (metadata shows actual creation date)",
            ),
            ("May 10, 2024", "Search warrants executed, Sarah Chen arrested"),
            ("May 10, 2024", "Marcus Rivera suspended from Port Authority"),
            ("May 12, 2024", "Digital forensics analysis reveals document tampering"),
            ("May 15, 2024", "Final investigation report compiled"),
        ]

        table = doc.add_table(rows=len(events) + 1, cols=2)
        table.style = "Light Grid Accent 1"

        table.rows[0].cells[0].text = "Date"
        table.rows[0].cells[1].text = "Event"
        for r in table.rows[0].cells:
            for p in r.paragraphs:
                for run in p.runs:
                    run.bold = True

        for i, (date, event) in enumerate(events, start=1):
            table.rows[i].cells[0].text = date
            table.rows[i].cells[1].text = event

        doc.add_paragraph("")
        doc.add_heading("Key Patterns Identified", 2)
        doc.add_paragraph(
            "1. STRUCTURED TRANSACTIONS: Multiple wire transfers just below $10K reporting threshold"
        )
        doc.add_paragraph(
            "2. DOCUMENT GAP: Two-week period (March 16-30) with no documentation"
        )
        doc.add_paragraph(
            "3. EVIDENCE TAMPERING: Backdated contract created after investigation began"
        )
        doc.add_paragraph(
            "4. COORDINATED ACTIVITY: Shell company registrations within days of each other"
        )

        doc.save(str(filepath))

    def gen_key_players(self):
        """Generate key players profile for narrative analysis."""
        filepath = self.output_dir / "Key_Players_Profile.docx"
        doc = Document()

        doc.add_heading("KEY PLAYERS DOSSIER", 0)
        doc.add_heading("Phantom Logistics Investigation", 2)

        players = [
            (
                "Sarah Chen",
                "CEO, Phantom Logistics LLC",
                """
- Age: 42, US Citizen (naturalized)
- Education: MBA, Stanford University
- Background: Former logistics executive at major shipping company
- Financial Status: Owns home in Malibu ($3.2M), luxury vehicles
- Criminal Record: None prior
- Role: Primary orchestrator of smuggling operation
- Key Relationships: Marcus Rivera (corrupt official), David Kim (former employee)
- Known Aliases: S. Chen, Dr. Chen (sometimes uses honorific despite no doctorate)
- Current Status: Detained pending trial, bail denied
""",
            ),
            (
                "Marcus Rivera",
                "Port Inspector, Port of Los Angeles",
                """
- Age: 38, US Citizen
- Employment: 12 years with Port Authority
- Badge Number: #4521
- Financial Status: Modest income, recent unexplained purchases
- Criminal Record: None prior
- Role: Insider facilitating container clearance in exchange for bribes
- Key Relationships: Sarah Chen (co-conspirator)
- Current Status: Suspended without pay, cooperating with investigation
""",
            ),
            (
                "David Kim",
                "Former Warehouse Manager, Phantom Logistics",
                """
- Age: 35, US Citizen
- Employment: Hired January 2024, resigned March 2024
- Background: 10 years in logistics industry
- Financial Status: Middle class, no unusual transactions
- Criminal Record: None
- Role: Whistleblower who exposed the operation
- Motivation: Ethical concerns about illegal activity
- Current Status: Protected witness, granted immunity
""",
            ),
            (
                "John Doe (fictitious)",
                "Listed CEO, Acme Imports Inc.",
                """
- Identity: UNKNOWN - likely fictitious
- No records matching this name at listed address
- Address: Virtual office in Wilmington, DE
- Role: Nominal head of shell company used for money transfers
- Assessment: Likely nominee or entirely fabricated identity
""",
            ),
        ]

        for name, title, bio in players:
            doc.add_heading(name, 2)
            doc.add_paragraph(f"Title: {title}").bold = True
            doc.add_paragraph(bio.strip())
            doc.add_paragraph("")

        doc.save(str(filepath))

    # ========== FINGERPRINT DETECTION Documents ==========

    def gen_invoice_template_v1(self):
        """Generate first invoice template (will be similar to v2)."""
        filepath = self.output_dir / "Invoice_Template_v1.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 16)
        c.drawString(inch, height - inch, "INVOICE")

        c.setFont("Helvetica", 10)
        c.drawString(inch, height - 1.5 * inch, "Acme Imports Inc.")
        c.drawString(inch, height - 1.7 * inch, "789 Trade Center Drive")
        c.drawString(inch, height - 1.9 * inch, "Los Angeles, CA 90012")

        c.drawString(inch, height - 2.5 * inch, "Invoice #: ACM-2024-0301")
        c.drawString(inch, height - 2.7 * inch, "Date: March 1, 2024")

        c.drawString(inch, height - 3.3 * inch, "BILL TO:")
        c.drawString(inch, height - 3.5 * inch, "Phantom Logistics LLC")
        c.drawString(
            inch, height - 3.7 * inch, "456 Commerce Blvd, Long Beach, CA 90802"
        )

        c.drawString(
            inch,
            height - 4.3 * inch,
            "Description: Professional Consulting Services - Q1 2024",
        )
        c.drawString(inch, height - 4.5 * inch, "  - Strategic Advisory")
        c.drawString(inch, height - 4.7 * inch, "  - Market Analysis")
        c.drawString(inch, height - 4.9 * inch, "  - Risk Management")

        c.setFont("Helvetica-Bold", 12)
        c.drawString(inch, height - 5.5 * inch, "TOTAL: $500,000.00")

        c.setFont("Helvetica", 9)
        c.drawString(
            inch,
            height - 6.5 * inch,
            "Wire transfer to: First National Bank, Account #1234567890",
        )

        c.save()

    def gen_invoice_template_v2(self):
        """Generate second invoice (near-duplicate of v1 with minor changes)."""
        filepath = self.output_dir / "Invoice_Template_v2.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        # Almost identical to v1, but different invoice number/amount
        c.setFont("Helvetica-Bold", 16)
        c.drawString(inch, height - inch, "INVOICE")

        c.setFont("Helvetica", 10)
        c.drawString(inch, height - 1.5 * inch, "Acme Imports Inc.")
        c.drawString(inch, height - 1.7 * inch, "789 Trade Center Drive")
        c.drawString(inch, height - 1.9 * inch, "Los Angeles, CA 90012")

        c.drawString(
            inch, height - 2.5 * inch, "Invoice #: ACM-2024-0315"
        )  # Different number
        c.drawString(
            inch, height - 2.7 * inch, "Date: March 15, 2024"
        )  # Different date

        c.drawString(inch, height - 3.3 * inch, "BILL TO:")
        c.drawString(inch, height - 3.5 * inch, "Phantom Logistics LLC")
        c.drawString(
            inch, height - 3.7 * inch, "456 Commerce Blvd, Long Beach, CA 90802"
        )

        c.drawString(
            inch,
            height - 4.3 * inch,
            "Description: Professional Consulting Services - Q1 2024",
        )
        c.drawString(inch, height - 4.5 * inch, "  - Strategic Advisory")
        c.drawString(inch, height - 4.7 * inch, "  - Market Analysis")
        c.drawString(inch, height - 4.9 * inch, "  - Risk Management")

        c.setFont("Helvetica-Bold", 12)
        c.drawString(
            inch, height - 5.5 * inch, "TOTAL: $750,000.00"
        )  # Different amount

        c.setFont("Helvetica", 9)
        c.drawString(
            inch,
            height - 6.5 * inch,
            "Wire transfer to: First National Bank, Account #1234567890",
        )

        c.save()

    # ========== SPECULATION / FACT COMPARISON Documents ==========

    def gen_witness_statement(self):
        """Generate witness statement with facts that need comparison."""
        filepath = self.output_dir / "Contradicting_Witness_Statement.pdf"
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width / 2, height - inch, "WITNESS STATEMENT")

        c.setFont("Helvetica", 10)
        c.drawString(
            inch, height - 1.5 * inch, "Witness: Maria Lopez (Administrative Assistant)"
        )
        c.drawString(
            inch,
            height - 1.7 * inch,
            f"Date: {self.date_final_report.strftime('%B %d, %Y')}",
        )

        content = """
STATEMENT:

I worked as Administrative Assistant for Phantom Logistics LLC from January 2024 
until the company was shut down in May.

OBSERVATIONS:

1. Sarah Chen was rarely in the office. She told me she had "meetings" but never 
   said where. She always left around 11am on Mondays.

2. I never saw Marcus Rivera at our office. Ms. Chen said she met him at "port 
   facilities" but I have no direct knowledge of these meetings.

3. The company received very few actual shipping clients. Most of our "business" 
   seemed to be wire transfers to various accounts.

4. David Kim, the warehouse manager, seemed uncomfortable. He kept asking questions 
   about why containers were being processed so quickly.

5. I was asked to backdate several documents. I refused and told Ms. Chen this was 
   against company policy. She said "just do it" and not to ask questions.

6. On May 8th, I saw Ms. Chen creating a contract on her computer. She asked me to 
   change the date in the footer to December 2023. I told her I was uncomfortable 
   with this. She did it herself.

7. The "consulting fees" on our invoices seemed excessive. $2.5 million for 
   "strategic advisory"? We never received any actual consulting reports.

I am providing this statement voluntarily and will testify if subpoenaed.

[SIGNATURE REDACTED FOR PROTECTION]
Maria Lopez
"""

        y = height - 2.2 * inch
        for line in content.strip().split("\n"):
            if y < 1.5 * inch:
                c.showPage()
                y = height - inch
                c.setFont("Helvetica", 10)
            c.drawString(inch, y, line)
            y -= 0.2 * inch

        c.save()

    def gen_missing_period(self):
        """Generate analysis of the suspicious March 16-30 gap."""
        filepath = self.output_dir / "Missing_Period_Analysis.txt"

        content = """INVESTIGATION ANALYSIS: THE MARCH 16-30 GAP
============================================
Prepared by: Senior Investigator J. Martinez
Date: May 14, 2024

BACKGROUND:
The two-week period between March 16 and March 30, 2024, represents a significant 
gap in the documentary evidence. This period falls immediately after the 
whistleblower complaint (March 15) and immediately before the falsified inspection 
report (April 1).

KNOWN EVENTS BEFORE GAP:
- March 10: Email from Rivera to Chen warning about "someone asking questions"
- March 15: Anonymous whistleblower complaint filed
- March 15: Last documented communication between Chen and Rivera in evidence

KNOWN EVENTS AFTER GAP:
- April 1: Inspector Rivera files inspection report marked "No Irregularities"
- April 14: Alleged cash payment (per handwritten note)
- April 15: $2.5M wire transfer to offshore account

QUESTIONS FOR SPECULATION MODE:
1. What happened during this two-week period?
2. Did Chen and Rivera meet in person to coordinate their response?
3. Were documents destroyed during this time?
4. Did the subjects become aware of the investigation?
5. Was anyone else brought into the conspiracy during this window?

INVESTIGATIVE LEADS:
- [ ] Subpoena cell phone records for Chen and Rivera (March 15-31)
- [ ] Check building access logs at Port facilities
- [ ] Review CCTV footage from known meeting locations
- [ ] Interview other Port Authority employees about Rivera's activities
- [ ] Check for cash withdrawals from Chen's accounts during this period

HYPOTHESIS:
This gap likely represents a "coordination window" during which the subjects:
1. Learned of the whistleblower complaint (possibly through internal leak)
2. Met to align their stories
3. Potentially destroyed incriminating evidence
4. Planned the falsified inspection report and backdated contract

RECOMMENDATION:
Focus interrogation efforts on this specific time period. The subjects' inability
to account for their activities during this window may reveal additional evidence
of conspiracy and obstruction.
"""

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

    def gen_extended_readme(self):
        """Update README with extended document information."""
        filepath = self.output_dir / "EXTENDED_FEATURES_README.md"

        content = f"""# Extended Tutorial Dataset

This directory contains the **extended tutorial dataset** that demonstrates ALL ArkhamMirror features, including advanced AI-powered investigation tools.

## 📋 Extended Documents (12 additional files)

### Contradiction Detection
| File | Purpose |
|------|---------|
| Deposition_Sarah_Chen.pdf | Chen's sworn testimony (contradicts Rivera) |
| Deposition_Marcus_Rivera.pdf | Rivera's testimony (contradicts Chen) |

Try: Navigate to **Contradictions** page and run detection. Should find conflicts about meeting frequency, cash payments, and knowledge of inspection bypasses.

### Red Flag Discovery
| File | Purpose |
|------|---------|
| Financial_Transactions_Q1_2024.txt | Transactions with structuring pattern (~$9,500 each) |
| Shell_Company_Registry.pdf | Offshore company information |

Try: Navigate to **Red Flags** page. Should detect:
- Structuring pattern (18+ transactions just below $10K)
- Round number transactions
- Multiple offshore entities

### Hidden Content / Metadata Forensics
| File | Purpose |
|------|---------|
| Backdated_Contract.pdf | Document with suspicious creation date |
| Suspicious_PDF_Analysis.txt | Forensics analysis of backdating |

Try: Navigate to **Metadata Forensics** page. Should detect creation/modification date anomalies.

### Big Picture / Narrative
| File | Purpose |
|------|---------|
| Investigation_Timeline_Summary.docx | Complete event chronology |
| Key_Players_Profile.docx | Entity profiles for analysis |

Try: Navigate to **Big Picture** page. Use "Generate Summary" to create executive briefing.

### Fingerprint Detection
| File | Purpose |
|------|---------|
| Invoice_Template_v1.pdf | First invoice (template) |
| Invoice_Template_v2.pdf | Near-duplicate invoice |

Try: Navigate to **Fingerprint** page. Should detect these as near-duplicates with high similarity score.

### Speculation / Fact Comparison
| File | Purpose |
|------|---------|
| Contradicting_Witness_Statement.pdf | Third-party witness account |
| Missing_Period_Analysis.txt | Analysis of March 16-30 gap |

Try: Navigate to **Speculation** page. Should generate investigative questions about the document gap.

## 🎯 Feature Testing Checklist

After ingesting all documents (15 basic + 12 extended = 27 total):

- [ ] **Contradiction Detection**: Find conflicts between Chen and Rivera depositions
- [ ] **Red Flag Discovery**: Detect structuring pattern in financial transactions
- [ ] **Hidden Content**: Find backdated document anomalies
- [ ] **Big Picture**: Generate executive summary of entire corpus
- [ ] **Narrative Reconstruction**: Build story from entity perspective
- [ ] **Speculation Mode**: Generate questions about March 16-30 gap
- [ ] **Fingerprint Detection**: Find Invoice Template v1/v2 as near-duplicates
- [ ] **Fact Comparison**: Compare claims across depositions and witness statements
- [ ] **Timeline Merging**: See comprehensive timeline with conflicts highlighted
- [ ] **Entity Influence**: View centrality metrics (Sarah Chen should be most central)

## 📊 Expected Results

After processing all 27 documents:

- **Entities**: ~25-30 (before deduplication)
- **Canonical Entities**: ~18-20 (after deduplication)
- **Contradictions Detected**: 3-5 (Chen vs Rivera statements)
- **Red Flags**: 15+ (structuring, offshore, gaps)
- **Fingerprint Matches**: 1-2 pairs (invoice templates)
- **Timeline Events**: 25+
- **Suspicious Gaps**: 1 major (March 16-30)

---
Generated: {datetime.now().strftime("%B %d, %Y")}
"""

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)


def main():
    """Main entry point."""
    # First run the base generator
    base_script = Path(__file__).parent / "generate_tutorial_dataset.py"
    if base_script.exists():
        print("[*] Running base tutorial dataset generator first...")
        import subprocess

        result = subprocess.run(
            [sys.executable, str(base_script)], capture_output=True, text=True
        )
        if result.returncode != 0:
            print(f"[WARNING] Base generator may have issues: {result.stderr}")

    # Then run extended generator
    generator = ExtendedTutorialGenerator()
    generator.generate_all()


if __name__ == "__main__":
    main()
