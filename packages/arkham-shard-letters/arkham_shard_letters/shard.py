"""
Letters Shard - Main Shard Implementation

Letter generation for ArkhamFrame - creates formal letters from templates
including FOIA requests, complaints, legal correspondence, and custom documents.
"""

import logging
import os
import re
import tempfile
from datetime import datetime
from typing import Any, Dict, List, Optional
from uuid import uuid4

from arkham_frame import ArkhamShard

from .models import (
    ExportFormat,
    Letter,
    LetterExportResult,
    LetterFilter,
    LetterStatistics,
    LetterStatus,
    LetterTemplate,
    LetterType,
    PlaceholderValue,
)

logger = logging.getLogger(__name__)


class LettersShard(ArkhamShard):
    """
    Letters Shard - Generates formal letters from templates.

    This shard provides:
    - Letter generation from templates with placeholder substitution
    - Multiple letter types (FOIA, complaint, demand, notice, etc.)
    - Template management with reusable content
    - Export to multiple formats (PDF, DOCX, HTML, Markdown, TXT)
    - Draft management and finalization workflow
    - Statistics and analytics
    """

    name = "letters"
    version = "0.1.0"
    description = "Letter generation - create formal letters from templates"

    def __init__(self):
        super().__init__()  # Auto-loads manifest from shard.yaml
        self.frame = None
        self._db = None
        self._events = None
        self._llm = None
        self._storage = None
        self._initialized = False
        self._letters_dir = None

    async def initialize(self, frame) -> None:
        """Initialize shard with frame services."""
        self.frame = frame
        self._db = frame.database
        self._events = frame.events
        self._llm = getattr(frame, "llm", None)
        self._storage = getattr(frame, "storage", None)

        # Setup letters output directory
        self._letters_dir = os.path.join(tempfile.gettempdir(), "arkham_letters")
        os.makedirs(self._letters_dir, exist_ok=True)

        # Create database schema
        await self._create_schema()

        # Register self in app state for API access
        if hasattr(frame, "app") and frame.app:
            frame.app.state.letters_shard = self

        self._initialized = True
        logger.info(f"LettersShard initialized (v{self.version})")

    async def shutdown(self) -> None:
        """Clean shutdown of shard."""
        self._initialized = False
        logger.info("LettersShard shutdown complete")

    def get_routes(self):
        """Return FastAPI router for this shard."""
        from .api import router
        return router

    # === Database Schema ===

    async def _create_schema(self) -> None:
        """Create database tables for letters shard."""
        if not self._db:
            logger.warning("Database not available, skipping schema creation")
            return

        # Letters table
        await self._db.execute("""
            CREATE TABLE IF NOT EXISTS arkham_letters (
                id TEXT PRIMARY KEY,
                title TEXT NOT NULL,
                letter_type TEXT NOT NULL,
                status TEXT DEFAULT 'draft',

                content TEXT DEFAULT '',
                template_id TEXT,

                recipient_name TEXT,
                recipient_address TEXT,
                recipient_email TEXT,
                sender_name TEXT,
                sender_address TEXT,
                sender_email TEXT,

                subject TEXT,
                reference_number TEXT,
                re_line TEXT,

                created_at TEXT,
                updated_at TEXT,
                finalized_at TEXT,
                sent_at TEXT,

                last_export_format TEXT,
                last_export_path TEXT,
                last_exported_at TEXT,

                metadata TEXT DEFAULT '{}'
            )
        """)

        # Templates table
        await self._db.execute("""
            CREATE TABLE IF NOT EXISTS arkham_letter_templates (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                letter_type TEXT NOT NULL,
                description TEXT,

                content_template TEXT NOT NULL,
                subject_template TEXT,

                placeholders TEXT DEFAULT '[]',
                required_placeholders TEXT DEFAULT '[]',

                default_sender_name TEXT,
                default_sender_address TEXT,
                default_sender_email TEXT,

                created_at TEXT,
                updated_at TEXT,
                metadata TEXT DEFAULT '{}'
            )
        """)

        # Create indexes
        await self._db.execute("""
            CREATE INDEX IF NOT EXISTS idx_letters_status ON arkham_letters(status)
        """)
        await self._db.execute("""
            CREATE INDEX IF NOT EXISTS idx_letters_type ON arkham_letters(letter_type)
        """)
        await self._db.execute("""
            CREATE INDEX IF NOT EXISTS idx_letters_created ON arkham_letters(created_at)
        """)
        await self._db.execute("""
            CREATE INDEX IF NOT EXISTS idx_letters_template ON arkham_letters(template_id)
        """)
        await self._db.execute("""
            CREATE INDEX IF NOT EXISTS idx_templates_type ON arkham_letter_templates(letter_type)
        """)

        logger.debug("Letters schema created/verified")

    # === Public API Methods - Letters ===

    async def create_letter(
        self,
        title: str,
        letter_type: LetterType,
        content: str = "",
        template_id: Optional[str] = None,
        recipient_name: Optional[str] = None,
        recipient_address: Optional[str] = None,
        subject: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Letter:
        """Create a new letter."""
        letter_id = str(uuid4())
        now = datetime.utcnow()

        letter = Letter(
            id=letter_id,
            title=title,
            letter_type=letter_type,
            status=LetterStatus.DRAFT,
            content=content,
            template_id=template_id,
            recipient_name=recipient_name,
            recipient_address=recipient_address,
            subject=subject,
            created_at=now,
            updated_at=now,
            metadata=metadata or {},
        )

        await self._save_letter(letter)

        # Emit event
        if self._events:
            await self._events.emit(
                "letters.letter.created",
                {
                    "letter_id": letter_id,
                    "letter_type": letter_type.value,
                    "title": title,
                    "template_id": template_id,
                },
                source=self.name,
            )

        return letter

    async def get_letter(self, letter_id: str) -> Optional[Letter]:
        """Get a letter by ID."""
        if not self._db:
            return None

        row = await self._db.fetch_one(
            "SELECT * FROM arkham_letters WHERE id = ?",
            [letter_id],
        )
        return self._row_to_letter(row) if row else None

    async def list_letters(
        self,
        filter: Optional[LetterFilter] = None,
        limit: int = 50,
        offset: int = 0,
    ) -> List[Letter]:
        """List letters with optional filtering."""
        if not self._db:
            return []

        query = "SELECT * FROM arkham_letters WHERE 1=1"
        params = []

        if filter:
            if filter.status:
                query += " AND status = ?"
                params.append(filter.status.value)
            if filter.letter_type:
                query += " AND letter_type = ?"
                params.append(filter.letter_type.value)
            if filter.template_id:
                query += " AND template_id = ?"
                params.append(filter.template_id)
            if filter.search_text:
                query += " AND (title LIKE ? OR content LIKE ? OR recipient_name LIKE ?)"
                search = f"%{filter.search_text}%"
                params.extend([search, search, search])

        query += " ORDER BY updated_at DESC LIMIT ? OFFSET ?"
        params.extend([limit, offset])

        rows = await self._db.fetch_all(query, params)
        return [self._row_to_letter(row) for row in rows]

    async def update_letter(
        self,
        letter_id: str,
        title: Optional[str] = None,
        content: Optional[str] = None,
        status: Optional[LetterStatus] = None,
        recipient_name: Optional[str] = None,
        recipient_address: Optional[str] = None,
        subject: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Optional[Letter]:
        """Update a letter."""
        letter = await self.get_letter(letter_id)
        if not letter:
            return None

        # Update fields
        if title is not None:
            letter.title = title
        if content is not None:
            letter.content = content
        if status is not None:
            old_status = letter.status
            letter.status = status
            # Track status transitions
            if status == LetterStatus.FINALIZED and old_status != LetterStatus.FINALIZED:
                letter.finalized_at = datetime.utcnow()
            elif status == LetterStatus.SENT and old_status != LetterStatus.SENT:
                letter.sent_at = datetime.utcnow()
        if recipient_name is not None:
            letter.recipient_name = recipient_name
        if recipient_address is not None:
            letter.recipient_address = recipient_address
        if subject is not None:
            letter.subject = subject
        if metadata is not None:
            letter.metadata.update(metadata)

        letter.updated_at = datetime.utcnow()

        await self._save_letter(letter, update=True)

        # Emit event
        if self._events:
            event_name = "letters.letter.updated"
            if status == LetterStatus.FINALIZED:
                event_name = "letters.letter.finalized"
            elif status == LetterStatus.SENT:
                event_name = "letters.letter.sent"

            await self._events.emit(
                event_name,
                {
                    "letter_id": letter_id,
                    "status": letter.status.value,
                    "title": letter.title,
                },
                source=self.name,
            )

        return letter

    async def delete_letter(self, letter_id: str) -> bool:
        """Delete a letter."""
        if not self._db:
            return False

        # Get letter to delete exported files
        letter = await self.get_letter(letter_id)
        if not letter:
            return False

        # Delete exported file if exists
        if letter.last_export_path and self._storage:
            try:
                await self._storage.delete(letter.last_export_path)
            except Exception as e:
                logger.warning(f"Failed to delete letter file: {e}")

        # Delete from database
        await self._db.execute(
            "DELETE FROM arkham_letters WHERE id = ?",
            [letter_id],
        )

        return True

    async def get_count(self, status: Optional[str] = None) -> int:
        """Get count of letters, optionally filtered by status."""
        if not self._db:
            return 0

        if status:
            result = await self._db.fetch_one(
                "SELECT COUNT(*) as count FROM arkham_letters WHERE status = ?",
                [status],
            )
        else:
            result = await self._db.fetch_one(
                "SELECT COUNT(*) as count FROM arkham_letters"
            )

        return result["count"] if result else 0

    # === Public API Methods - Templates ===

    async def create_template(
        self,
        name: str,
        letter_type: LetterType,
        description: str,
        content_template: str,
        subject_template: Optional[str] = None,
        default_sender_name: Optional[str] = None,
        default_sender_address: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> LetterTemplate:
        """Create a new letter template."""
        template_id = str(uuid4())
        now = datetime.utcnow()

        # Extract placeholders from template
        placeholders = self._extract_placeholders(content_template)
        if subject_template:
            placeholders.update(self._extract_placeholders(subject_template))

        template = LetterTemplate(
            id=template_id,
            name=name,
            letter_type=letter_type,
            description=description,
            content_template=content_template,
            subject_template=subject_template,
            placeholders=list(placeholders),
            required_placeholders=[],  # Would be configured separately
            default_sender_name=default_sender_name,
            default_sender_address=default_sender_address,
            created_at=now,
            updated_at=now,
            metadata=metadata or {},
        )

        await self._save_template(template)

        # Emit event
        if self._events:
            await self._events.emit(
                "letters.template.created",
                {
                    "template_id": template_id,
                    "name": name,
                    "letter_type": letter_type.value,
                    "placeholders": template.placeholders,
                },
                source=self.name,
            )

        return template

    async def get_template(self, template_id: str) -> Optional[LetterTemplate]:
        """Get a template by ID."""
        if not self._db:
            return None

        row = await self._db.fetch_one(
            "SELECT * FROM arkham_letter_templates WHERE id = ?",
            [template_id],
        )
        return self._row_to_template(row) if row else None

    async def list_templates(
        self,
        letter_type: Optional[LetterType] = None,
        limit: int = 100,
        offset: int = 0,
    ) -> List[LetterTemplate]:
        """List letter templates."""
        if not self._db:
            return []

        query = "SELECT * FROM arkham_letter_templates WHERE 1=1"
        params = []

        if letter_type:
            query += " AND letter_type = ?"
            params.append(letter_type.value)

        query += " ORDER BY name LIMIT ? OFFSET ?"
        params.extend([limit, offset])

        rows = await self._db.fetch_all(query, params)
        return [self._row_to_template(row) for row in rows]

    async def update_template(
        self,
        template_id: str,
        name: Optional[str] = None,
        description: Optional[str] = None,
        content_template: Optional[str] = None,
        subject_template: Optional[str] = None,
    ) -> Optional[LetterTemplate]:
        """Update a template."""
        template = await self.get_template(template_id)
        if not template:
            return None

        if name is not None:
            template.name = name
        if description is not None:
            template.description = description
        if content_template is not None:
            template.content_template = content_template
            # Re-extract placeholders
            template.placeholders = list(self._extract_placeholders(content_template))
        if subject_template is not None:
            template.subject_template = subject_template

        template.updated_at = datetime.utcnow()

        await self._save_template(template, update=True)

        # Emit event
        if self._events:
            await self._events.emit(
                "letters.template.updated",
                {
                    "template_id": template_id,
                    "name": template.name,
                },
                source=self.name,
            )

        return template

    async def delete_template(self, template_id: str) -> bool:
        """Delete a template."""
        if not self._db:
            return False

        await self._db.execute(
            "DELETE FROM arkham_letter_templates WHERE id = ?",
            [template_id],
        )
        return True

    # === Template Application ===

    async def apply_template(
        self,
        template_id: str,
        title: str,
        placeholder_values: List[PlaceholderValue],
        recipient_name: Optional[str] = None,
        recipient_address: Optional[str] = None,
    ) -> Letter:
        """Create a new letter from a template with placeholder substitution."""
        template = await self.get_template(template_id)
        if not template:
            raise ValueError(f"Template {template_id} not found")

        # Build placeholder map
        placeholder_map = {pv.key: pv.value for pv in placeholder_values}

        # Render content
        content = self._render_template(template.content_template, placeholder_map)

        # Render subject if template has one
        subject = None
        if template.subject_template:
            subject = self._render_template(template.subject_template, placeholder_map)

        # Create letter with rendered content
        letter = await self.create_letter(
            title=title,
            letter_type=template.letter_type,
            content=content,
            template_id=template_id,
            recipient_name=recipient_name or "",
            recipient_address=recipient_address or "",
            subject=subject,
            metadata={
                "from_template": template_id,
                "template_name": template.name,
                "placeholders_used": list(placeholder_map.keys()),
            },
        )

        # Emit event
        if self._events:
            await self._events.emit(
                "letters.template.applied",
                {
                    "letter_id": letter.id,
                    "template_id": template_id,
                    "template_name": template.name,
                },
                source=self.name,
            )

        return letter

    # === Export Methods ===

    async def export_letter(
        self,
        letter_id: str,
        export_format: ExportFormat,
    ) -> LetterExportResult:
        """Export a letter to a file format."""
        import time
        start_time = time.time()

        letter = await self.get_letter(letter_id)
        if not letter:
            return LetterExportResult(
                letter_id=letter_id,
                success=False,
                export_format=export_format,
                errors=["Letter not found"],
            )

        try:
            # Generate export content
            file_content = await self._generate_export(letter, export_format)

            # Save to letters directory
            file_name = f"letter_{letter_id}.{export_format.value}"
            file_path = os.path.join(self._letters_dir, file_name)

            # Write file
            with open(file_path, 'wb') as f:
                f.write(file_content)

            # Update letter with export info
            letter.last_export_format = export_format
            letter.last_export_path = file_path
            letter.last_exported_at = datetime.utcnow()
            await self._save_letter(letter, update=True)

            processing_time = (time.time() - start_time) * 1000

            # Emit event
            if self._events:
                await self._events.emit(
                    "letters.letter.exported",
                    {
                        "letter_id": letter_id,
                        "export_format": export_format.value,
                        "file_path": file_path,
                    },
                    source=self.name,
                )

            return LetterExportResult(
                letter_id=letter_id,
                success=True,
                export_format=export_format,
                file_path=file_path,
                file_size=len(file_content),
                processing_time_ms=processing_time,
            )

        except Exception as e:
            logger.error(f"Letter export failed: {e}")
            return LetterExportResult(
                letter_id=letter_id,
                success=False,
                export_format=export_format,
                errors=[str(e)],
            )

    # === Statistics ===

    async def get_statistics(self) -> LetterStatistics:
        """Get statistics about letters in the system."""
        if not self._db:
            return LetterStatistics()

        # Total letters
        total = await self._db.fetch_one(
            "SELECT COUNT(*) as count FROM arkham_letters"
        )
        total_letters = total["count"] if total else 0

        # By status
        status_rows = await self._db.fetch_all(
            "SELECT status, COUNT(*) as count FROM arkham_letters GROUP BY status"
        )
        by_status = {row["status"]: row["count"] for row in status_rows}

        # By type
        type_rows = await self._db.fetch_all(
            "SELECT letter_type, COUNT(*) as count FROM arkham_letters GROUP BY letter_type"
        )
        by_type = {row["letter_type"]: row["count"] for row in type_rows}

        # Templates
        templates = await self._db.fetch_one(
            "SELECT COUNT(*) as count FROM arkham_letter_templates"
        )
        total_templates = templates["count"] if templates else 0

        # Template types
        template_type_rows = await self._db.fetch_all(
            "SELECT letter_type, COUNT(*) as count FROM arkham_letter_templates GROUP BY letter_type"
        )
        by_template_type = {row["letter_type"]: row["count"] for row in template_type_rows}

        # Recent letters (stub - would need date filtering)
        letters_last_24h = 0
        letters_last_7d = 0
        letters_last_30d = 0

        # Exports (stub - would track in separate table)
        total_exports = 0
        by_export_format = {}

        return LetterStatistics(
            total_letters=total_letters,
            by_status=by_status,
            by_type=by_type,
            total_templates=total_templates,
            by_template_type=by_template_type,
            letters_last_24h=letters_last_24h,
            letters_last_7d=letters_last_7d,
            letters_last_30d=letters_last_30d,
            total_exports=total_exports,
            by_export_format=by_export_format,
        )

    # === Private Helper Methods ===

    def _extract_placeholders(self, template: str) -> set:
        """Extract placeholder variables from template string."""
        # Find all {{placeholder}} patterns
        pattern = r'\{\{(\w+)\}\}'
        matches = re.findall(pattern, template)
        return set(matches)

    def _render_template(self, template: str, placeholder_map: Dict[str, str]) -> str:
        """Render a template by substituting placeholders."""
        result = template
        for key, value in placeholder_map.items():
            placeholder = f"{{{{{key}}}}}"  # {{key}}
            result = result.replace(placeholder, value)
        return result

    async def _generate_export(self, letter: Letter, export_format: ExportFormat) -> bytes:
        """Generate export file content."""
        # Build letter components
        letter_data = self._build_letter_data(letter)
        full_text = self._format_letter_text(letter_data)

        # Export based on format
        if export_format == ExportFormat.TXT:
            return full_text.encode('utf-8')
        elif export_format == ExportFormat.MARKDOWN:
            return self._generate_markdown(letter, letter_data)
        elif export_format == ExportFormat.HTML:
            return self._generate_html(letter, letter_data)
        elif export_format == ExportFormat.PDF:
            return self._generate_pdf(letter, letter_data)
        elif export_format == ExportFormat.DOCX:
            return self._generate_docx(letter, letter_data)
        else:
            return full_text.encode('utf-8')

    def _build_letter_data(self, letter: Letter) -> Dict[str, Any]:
        """Build structured letter data for rendering."""
        return {
            "sender_name": letter.sender_name or "",
            "sender_address": letter.sender_address or "",
            "sender_email": letter.sender_email or "",
            "recipient_name": letter.recipient_name or "",
            "recipient_address": letter.recipient_address or "",
            "recipient_email": letter.recipient_email or "",
            "date": datetime.utcnow().strftime("%B %d, %Y"),
            "subject": letter.subject or "",
            "re_line": letter.re_line or "",
            "reference_number": letter.reference_number or "",
            "content": letter.content or "",
            "title": letter.title,
            "letter_type": letter.letter_type.value,
        }

    def _format_letter_text(self, data: Dict[str, Any]) -> str:
        """Format letter as plain text."""
        lines = []

        # Header
        if data["sender_name"]:
            lines.append(data["sender_name"])
        if data["sender_address"]:
            lines.append(data["sender_address"])
        if data["sender_email"]:
            lines.append(data["sender_email"])

        lines.append("")
        lines.append(data["date"])
        lines.append("")

        # Recipient
        if data["recipient_name"]:
            lines.append(data["recipient_name"])
        if data["recipient_address"]:
            lines.append(data["recipient_address"])

        lines.append("")

        # Subject/RE
        if data["subject"]:
            lines.append(f"Subject: {data['subject']}")
        if data["re_line"]:
            lines.append(f"RE: {data['re_line']}")
        if data["reference_number"]:
            lines.append(f"Reference: {data['reference_number']}")

        if data["subject"] or data["re_line"] or data["reference_number"]:
            lines.append("")

        # Salutation
        if data["recipient_name"]:
            lines.append(f"Dear {data['recipient_name']}:")
        else:
            lines.append("Dear Sir/Madam:")

        lines.append("")

        # Body
        lines.append(data["content"])

        lines.append("")

        # Closing
        lines.append("Sincerely,")
        lines.append("")
        if data["sender_name"]:
            lines.append(data["sender_name"])

        return "\n".join(lines)

    def _generate_markdown(self, letter: Letter, data: Dict[str, Any]) -> bytes:
        """Generate Markdown formatted letter."""
        md = f"""# {data['title']}

**Date:** {data['date']}
**Type:** {data['letter_type'].replace('_', ' ').title()}

---

"""
        if data["sender_name"]:
            md += f"**From:** {data['sender_name']}\n"
        if data["sender_address"]:
            md += f"{data['sender_address']}\n"

        md += "\n"

        if data["recipient_name"]:
            md += f"**To:** {data['recipient_name']}\n"
        if data["recipient_address"]:
            md += f"{data['recipient_address']}\n"

        md += "\n---\n\n"

        if data["subject"]:
            md += f"**Subject:** {data['subject']}\n\n"
        if data["re_line"]:
            md += f"**RE:** {data['re_line']}\n\n"
        if data["reference_number"]:
            md += f"**Reference:** {data['reference_number']}\n\n"

        if data["recipient_name"]:
            md += f"Dear {data['recipient_name']},\n\n"
        else:
            md += "Dear Sir/Madam,\n\n"

        md += data["content"]
        md += "\n\nSincerely,\n\n"
        if data["sender_name"]:
            md += data["sender_name"]

        return md.encode('utf-8')

    def _generate_html(self, letter: Letter, data: Dict[str, Any]) -> bytes:
        """Generate HTML formatted letter."""
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{data['title']}</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
        .header {{ margin-bottom: 30px; }}
        .date {{ margin: 20px 0; }}
        .recipient {{ margin-bottom: 20px; }}
        .subject {{ font-weight: bold; margin-bottom: 20px; }}
        .salutation {{ margin-bottom: 15px; }}
        .body {{ line-height: 1.6; margin-bottom: 30px; white-space: pre-wrap; }}
        .closing {{ margin-top: 30px; }}
        .signature {{ margin-top: 40px; }}
    </style>
</head>
<body>
    <div class="header">
"""
        if data["sender_name"]:
            html += f"        <div>{data['sender_name']}</div>\n"
        if data["sender_address"]:
            html += f"        <div>{data['sender_address']}</div>\n"
        if data["sender_email"]:
            html += f"        <div>{data['sender_email']}</div>\n"

        html += f"""    </div>
    <div class="date">{data['date']}</div>
    <div class="recipient">
"""
        if data["recipient_name"]:
            html += f"        <div>{data['recipient_name']}</div>\n"
        if data["recipient_address"]:
            html += f"        <div>{data['recipient_address']}</div>\n"

        html += "    </div>\n"

        if data["subject"]:
            html += f"    <div class='subject'>Subject: {data['subject']}</div>\n"
        if data["re_line"]:
            html += f"    <div class='subject'>RE: {data['re_line']}</div>\n"
        if data["reference_number"]:
            html += f"    <div>Reference: {data['reference_number']}</div>\n"

        salutation = f"Dear {data['recipient_name']}:" if data["recipient_name"] else "Dear Sir/Madam:"
        html += f"""    <div class="salutation">{salutation}</div>
    <div class="body">{data['content']}</div>
    <div class="closing">Sincerely,</div>
"""
        if data["sender_name"]:
            html += f"    <div class='signature'>{data['sender_name']}</div>\n"

        html += "</body>\n</html>"

        return html.encode('utf-8')

    def _generate_pdf(self, letter: Letter, data: Dict[str, Any]) -> bytes:
        """Generate PDF formatted letter using reportlab."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter as LETTER_SIZE
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from io import BytesIO
        except ImportError:
            logger.error("reportlab not installed - returning text fallback")
            return self._format_letter_text(data).encode('utf-8')

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=LETTER_SIZE,
                                leftMargin=1*inch, rightMargin=1*inch,
                                topMargin=1*inch, bottomMargin=1*inch)
        styles = getSampleStyleSheet()
        story = []

        # Define styles
        normal_style = ParagraphStyle(
            'LetterNormal',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=12,
            leading=16,
        )
        bold_style = ParagraphStyle(
            'LetterBold',
            parent=normal_style,
            fontName='Times-Bold',
        )

        # Sender header
        if data["sender_name"]:
            story.append(Paragraph(data["sender_name"], bold_style))
        if data["sender_address"]:
            for line in data["sender_address"].split('\n'):
                story.append(Paragraph(line, normal_style))
        if data["sender_email"]:
            story.append(Paragraph(data["sender_email"], normal_style))

        story.append(Spacer(1, 24))

        # Date
        story.append(Paragraph(data["date"], normal_style))
        story.append(Spacer(1, 24))

        # Recipient
        if data["recipient_name"]:
            story.append(Paragraph(data["recipient_name"], normal_style))
        if data["recipient_address"]:
            for line in data["recipient_address"].split('\n'):
                story.append(Paragraph(line, normal_style))

        story.append(Spacer(1, 24))

        # Subject/RE
        if data["subject"]:
            story.append(Paragraph(f"<b>Subject:</b> {data['subject']}", normal_style))
        if data["re_line"]:
            story.append(Paragraph(f"<b>RE:</b> {data['re_line']}", normal_style))
        if data["reference_number"]:
            story.append(Paragraph(f"<b>Reference:</b> {data['reference_number']}", normal_style))

        if data["subject"] or data["re_line"] or data["reference_number"]:
            story.append(Spacer(1, 18))

        # Salutation
        salutation = f"Dear {data['recipient_name']}:" if data["recipient_name"] else "Dear Sir/Madam:"
        story.append(Paragraph(salutation, normal_style))
        story.append(Spacer(1, 12))

        # Body - handle paragraphs
        for para in data["content"].split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.replace('\n', '<br/>'), normal_style))
                story.append(Spacer(1, 12))

        # Closing
        story.append(Spacer(1, 24))
        story.append(Paragraph("Sincerely,", normal_style))
        story.append(Spacer(1, 36))
        if data["sender_name"]:
            story.append(Paragraph(data["sender_name"], normal_style))

        doc.build(story)
        return buffer.getvalue()

    def _generate_docx(self, letter: Letter, data: Dict[str, Any]) -> bytes:
        """Generate DOCX formatted letter using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
        except ImportError:
            logger.error("python-docx not installed - returning text fallback")
            return self._format_letter_text(data).encode('utf-8')

        document = Document()

        # Set margins
        for section in document.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        # Sender header
        if data["sender_name"]:
            p = document.add_paragraph()
            run = p.add_run(data["sender_name"])
            run.bold = True
            run.font.size = Pt(12)

        if data["sender_address"]:
            for line in data["sender_address"].split('\n'):
                p = document.add_paragraph(line)
                p.paragraph_format.space_after = Pt(0)

        if data["sender_email"]:
            p = document.add_paragraph(data["sender_email"])
            p.paragraph_format.space_after = Pt(0)

        # Blank line
        document.add_paragraph()

        # Date
        document.add_paragraph(data["date"])

        # Blank line
        document.add_paragraph()

        # Recipient
        if data["recipient_name"]:
            document.add_paragraph(data["recipient_name"])
        if data["recipient_address"]:
            for line in data["recipient_address"].split('\n'):
                p = document.add_paragraph(line)
                p.paragraph_format.space_after = Pt(0)

        # Blank line
        document.add_paragraph()

        # Subject/RE
        if data["subject"]:
            p = document.add_paragraph()
            p.add_run("Subject: ").bold = True
            p.add_run(data["subject"])

        if data["re_line"]:
            p = document.add_paragraph()
            p.add_run("RE: ").bold = True
            p.add_run(data["re_line"])

        if data["reference_number"]:
            p = document.add_paragraph()
            p.add_run("Reference: ").bold = True
            p.add_run(data["reference_number"])

        if data["subject"] or data["re_line"] or data["reference_number"]:
            document.add_paragraph()

        # Salutation
        salutation = f"Dear {data['recipient_name']}:" if data["recipient_name"] else "Dear Sir/Madam:"
        document.add_paragraph(salutation)

        # Blank line
        document.add_paragraph()

        # Body
        for para in data["content"].split('\n\n'):
            if para.strip():
                p = document.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(12)

        # Closing
        document.add_paragraph()
        document.add_paragraph("Sincerely,")
        document.add_paragraph()
        document.add_paragraph()

        if data["sender_name"]:
            document.add_paragraph(data["sender_name"])

        # Save to buffer
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    async def _save_letter(self, letter: Letter, update: bool = False) -> None:
        """Save a letter to the database."""
        if not self._db:
            return

        import json
        data = (
            letter.id,
            letter.title,
            letter.letter_type.value,
            letter.status.value,
            letter.content,
            letter.template_id,
            letter.recipient_name,
            letter.recipient_address,
            letter.recipient_email,
            letter.sender_name,
            letter.sender_address,
            letter.sender_email,
            letter.subject,
            letter.reference_number,
            letter.re_line,
            letter.created_at.isoformat(),
            letter.updated_at.isoformat(),
            letter.finalized_at.isoformat() if letter.finalized_at else None,
            letter.sent_at.isoformat() if letter.sent_at else None,
            letter.last_export_format.value if letter.last_export_format else None,
            letter.last_export_path,
            letter.last_exported_at.isoformat() if letter.last_exported_at else None,
            json.dumps(letter.metadata),
        )

        if update:
            await self._db.execute("""
                UPDATE arkham_letters SET
                    title=?, letter_type=?, status=?, content=?, template_id=?,
                    recipient_name=?, recipient_address=?, recipient_email=?,
                    sender_name=?, sender_address=?, sender_email=?,
                    subject=?, reference_number=?, re_line=?,
                    created_at=?, updated_at=?, finalized_at=?, sent_at=?,
                    last_export_format=?, last_export_path=?, last_exported_at=?,
                    metadata=?
                WHERE id=?
            """, data[1:] + (letter.id,))
        else:
            await self._db.execute("""
                INSERT INTO arkham_letters (
                    id, title, letter_type, status, content, template_id,
                    recipient_name, recipient_address, recipient_email,
                    sender_name, sender_address, sender_email,
                    subject, reference_number, re_line,
                    created_at, updated_at, finalized_at, sent_at,
                    last_export_format, last_export_path, last_exported_at,
                    metadata
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, data)

    async def _save_template(self, template: LetterTemplate, update: bool = False) -> None:
        """Save a template to the database."""
        if not self._db:
            return

        import json
        data = (
            template.id,
            template.name,
            template.letter_type.value,
            template.description,
            template.content_template,
            template.subject_template,
            json.dumps(template.placeholders),
            json.dumps(template.required_placeholders),
            template.default_sender_name,
            template.default_sender_address,
            template.default_sender_email,
            template.created_at.isoformat(),
            template.updated_at.isoformat(),
            json.dumps(template.metadata),
        )

        if update:
            await self._db.execute("""
                UPDATE arkham_letter_templates SET
                    name=?, letter_type=?, description=?,
                    content_template=?, subject_template=?,
                    placeholders=?, required_placeholders=?,
                    default_sender_name=?, default_sender_address=?, default_sender_email=?,
                    created_at=?, updated_at=?, metadata=?
                WHERE id=?
            """, data[1:] + (template.id,))
        else:
            await self._db.execute("""
                INSERT INTO arkham_letter_templates (
                    id, name, letter_type, description,
                    content_template, subject_template,
                    placeholders, required_placeholders,
                    default_sender_name, default_sender_address, default_sender_email,
                    created_at, updated_at, metadata
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, data)

    def _row_to_letter(self, row: Dict[str, Any]) -> Letter:
        """Convert database row to Letter object."""
        import json
        return Letter(
            id=row["id"],
            title=row["title"],
            letter_type=LetterType(row["letter_type"]),
            status=LetterStatus(row["status"]),
            content=row["content"] or "",
            template_id=row["template_id"],
            recipient_name=row["recipient_name"],
            recipient_address=row["recipient_address"],
            recipient_email=row["recipient_email"],
            sender_name=row["sender_name"],
            sender_address=row["sender_address"],
            sender_email=row["sender_email"],
            subject=row["subject"],
            reference_number=row["reference_number"],
            re_line=row["re_line"],
            created_at=datetime.fromisoformat(row["created_at"]) if row["created_at"] else datetime.utcnow(),
            updated_at=datetime.fromisoformat(row["updated_at"]) if row["updated_at"] else datetime.utcnow(),
            finalized_at=datetime.fromisoformat(row["finalized_at"]) if row["finalized_at"] else None,
            sent_at=datetime.fromisoformat(row["sent_at"]) if row["sent_at"] else None,
            last_export_format=ExportFormat(row["last_export_format"]) if row["last_export_format"] else None,
            last_export_path=row["last_export_path"],
            last_exported_at=datetime.fromisoformat(row["last_exported_at"]) if row["last_exported_at"] else None,
            metadata=json.loads(row["metadata"] or "{}"),
        )

    def _row_to_template(self, row: Dict[str, Any]) -> LetterTemplate:
        """Convert database row to LetterTemplate object."""
        import json
        return LetterTemplate(
            id=row["id"],
            name=row["name"],
            letter_type=LetterType(row["letter_type"]),
            description=row["description"] or "",
            content_template=row["content_template"],
            subject_template=row["subject_template"],
            placeholders=json.loads(row["placeholders"] or "[]"),
            required_placeholders=json.loads(row["required_placeholders"] or "[]"),
            default_sender_name=row["default_sender_name"],
            default_sender_address=row["default_sender_address"],
            default_sender_email=row["default_sender_email"],
            created_at=datetime.fromisoformat(row["created_at"]) if row["created_at"] else datetime.utcnow(),
            updated_at=datetime.fromisoformat(row["updated_at"]) if row["updated_at"] else datetime.utcnow(),
            metadata=json.loads(row["metadata"] or "{}"),
        )
