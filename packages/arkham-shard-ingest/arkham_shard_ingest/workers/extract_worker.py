"""
ExtractWorker - Text and metadata extraction from documents.

Extracts text and metadata from PDF, DOCX, XLSX, CSV, TXT, EML using CPU-based
libraries. Prefers EXIFTool for metadata when available; falls back to
format-specific extractors. Single source of truth for document metadata.
Part of the cpu-extract worker pool for document processing.
"""

import asyncio
import logging
import os
from pathlib import Path
from typing import Dict, Any, Optional

from arkham_frame.workers.base import BaseWorker

from .metadata_extraction import (
    run_exiftool,
    normalize_exiftool_to_metadata,
    add_filesystem_times,
)

logger = logging.getLogger(__name__)

# MIME type -> internal file_type for dispatch (must match _extract_* methods)
MIME_TO_FILE_TYPE = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
    "application/vnd.oasis.opendocument.text": "odt",
    "application/rtf": "txt",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xlsx",
    "text/csv": "csv",
    "text/tab-separated-values": "csv",
    "text/plain": "txt",
    "text/markdown": "txt",
    "message/rfc822": "eml",
    "application/vnd.ms-outlook": "eml",
}
# Supported file types for extraction (must match _extract_* methods)
SUPPORTED_FILE_TYPES = ("pdf", "docx", "xlsx", "csv", "txt", "eml")

# Fallback: extension -> file_type when mime not in map
EXT_TO_FILE_TYPE = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".xlsx": "xlsx",
    ".xls": "xlsx",
    ".csv": "csv",
    ".tsv": "csv",
    ".txt": "txt",
    ".md": "txt",
    ".log": "txt",
    ".eml": "eml",
    ".emlx": "eml",
}


class ExtractWorker(BaseWorker):
    """
    Worker for extracting text from documents.

    Supports PDF, DOCX, and XLSX file formats. Handles various edge cases
    including missing files, corrupted documents, password-protected files,
    and missing dependencies.

    Uses the cpu-extract pool for document text extraction tasks.
    """

    pool = "cpu-extract"
    name = "ExtractWorker"

    # Configuration
    poll_interval = 1.0
    heartbeat_interval = 10.0
    idle_timeout = 300.0  # 5 minutes
    job_timeout = 120.0   # 2 minutes for large files
    max_retries = 2

    def __init__(self, *args, **kwargs):
        """Initialize worker and check for required dependencies."""
        super().__init__(*args, **kwargs)
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which extraction libraries are available."""
        self._has_pypdf = False
        self._has_docx = False
        self._has_openpyxl = False

        try:
            import pypdf
            self._has_pypdf = True
        except ImportError:
            logger.warning("pypdf not installed - PDF extraction unavailable")

        try:
            import docx
            self._has_docx = True
        except ImportError:
            logger.warning("python-docx not installed - DOCX extraction unavailable")

        try:
            import openpyxl
            self._has_openpyxl = True
        except ImportError:
            logger.warning("openpyxl not installed - XLSX extraction unavailable")

        if not any([self._has_pypdf, self._has_docx, self._has_openpyxl]):
            logger.error(
                "No extraction libraries available! "
                "Install pypdf, python-docx, and/or openpyxl"
            )

    def _resolve_path(self, file_path: str) -> Path:
        """
        Resolve file path using DATA_SILO_PATH for Docker/portable deployments.

        Args:
            file_path: Path from payload (may be relative or absolute)

        Returns:
            Resolved absolute Path
        """
        if not os.path.isabs(file_path):
            data_silo = os.environ.get("DATA_SILO_PATH", ".")
            return Path(data_silo) / file_path
        return Path(file_path)

    async def process_job(self, job_id: str, payload: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract text and metadata from a document.

        Payload:
            file_path: Path to the file to extract from (required)
            file_type: Optional - "pdf", "docx", "xlsx", "csv", "txt", "eml"
            file_info: Optional dict with mime_type (used for routing when file_type absent)
            mime_type: Optional - used for routing when file_info.mime_type not present

        Returns:
            dict with success, text, pages, document_metadata, file_path, file_type.
        """
        file_path = payload.get("file_path")
        if not file_path:
            raise ValueError("Missing required parameter: file_path")

        file_path = str(self._resolve_path(file_path))
        path = Path(file_path)

        # Derive file_type from mime_type (single source of type) then extension
        file_info = payload.get("file_info") or {}
        mime_type = file_info.get("mime_type") or payload.get("mime_type") or ""
        file_type = payload.get("file_type", "").lower()
        if not file_type and mime_type:
            file_type = MIME_TO_FILE_TYPE.get(mime_type.strip(), "")
        if not file_type:
            ext = path.suffix.lower()
            file_type = EXT_TO_FILE_TYPE.get(ext, "")

        if not file_type:
            raise ValueError(
                "Could not determine file type. Provide file_type, mime_type, or use a supported extension: "
                + ", ".join(SUPPORTED_FILE_TYPES)
            )

        if file_type not in SUPPORTED_FILE_TYPES:
            # Map odt/rtf to txt for simple read; otherwise reject
            if file_type in ("odt", "rtf"):
                file_type = "txt"
            else:
                raise ValueError(
                    f"Unsupported file_type: {file_type}. Supported: {', '.join(SUPPORTED_FILE_TYPES)}"
                )

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        if not path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")

        logger.info(
            "ExtractWorker processing %s file: %s (job %s)",
            file_type.upper(),
            path.name,
            job_id,
        )

        try:
            if file_type == "pdf":
                result = await self._extract_pdf(path)
            elif file_type == "docx":
                result = await self._extract_docx(path)
            elif file_type == "xlsx":
                result = await self._extract_xlsx(path)
            elif file_type == "txt":
                result = await self._extract_text(path)
            elif file_type == "csv":
                result = await self._extract_csv(path)
            elif file_type == "eml":
                result = await self._extract_eml(path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Ensure document_metadata dict exists
            doc_meta = result.get("document_metadata") or {}
            result["document_metadata"] = doc_meta

            # EXIFTool-first: run when available and merge into document_metadata
            raw_exif = run_exiftool(path)
            if raw_exif:
                exif_meta = normalize_exiftool_to_metadata(raw_exif)
                for k, v in exif_meta.items():
                    if k == "exiftool_metadata":
                        doc_meta["exiftool_metadata"] = v
                    elif v and (k not in doc_meta or not doc_meta[k]):
                        doc_meta[k] = v
                logger.debug("Merged EXIFTool metadata for job %s", job_id)

            # Filesystem times (single source of truth)
            add_filesystem_times(path, doc_meta)

            # num_pages from extraction result if not set
            if "num_pages" not in doc_meta or doc_meta["num_pages"] is None:
                doc_meta["num_pages"] = result.get("pages", 0)

            # PII detection is done in ingest shard _register_document (PII shard or fallback)

            result["file_path"] = str(file_path)
            result["file_type"] = file_type
            result["success"] = True

            logger.info(
                "ExtractWorker completed %s: %s pages, %s chars",
                file_type.upper(),
                result.get("pages", 0),
                len(result.get("text", "")),
            )
            return result

        except Exception as e:
            error_msg = f"Extraction failed for {file_type.upper()}: {str(e)}"
            logger.error("ExtractWorker error (job %s): %s", job_id, error_msg)
            return {
                "success": False,
                "text": "",
                "pages": 0,
                "document_metadata": {},
                "error": error_msg,
                "file_path": str(file_path),
                "file_type": file_type,
            }

    async def _extract_pdf(self, path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from PDF file.

        Args:
            path: Path to PDF file

        Returns:
            dict with text, page count, and document_metadata

        Raises:
            ImportError: If pypdf is not installed
            Exception: For PDF reading errors
        """
        if not self._has_pypdf:
            raise ImportError(
                "pypdf library not installed. "
                "Install with: pip install pypdf"
            )

        from pypdf import PdfReader

        # Run in executor to avoid blocking
        def extract():
            try:
                reader = PdfReader(str(path))

                # Check for encryption (store in metadata for single source of truth)
                is_encrypted = getattr(reader, "is_encrypted", False) or False
                document_metadata = {"is_encrypted": is_encrypted}
                if is_encrypted:
                    raise ValueError(
                        "PDF is password-protected. "
                        "Encrypted PDFs are not supported."
                    )

                # Extract PDF metadata (author, title, creator, etc.)
                if reader.metadata:
                    meta = reader.metadata
                    # Standard PDF metadata fields
                    if meta.author:
                        document_metadata["author"] = str(meta.author)
                    if meta.title:
                        document_metadata["title"] = str(meta.title)
                    if meta.subject:
                        document_metadata["subject"] = str(meta.subject)
                    if meta.creator:
                        document_metadata["creator"] = str(meta.creator)
                    if meta.producer:
                        document_metadata["producer"] = str(meta.producer)
                    if meta.creation_date:
                        document_metadata["creation_date"] = (
                            meta.creation_date.isoformat()
                            if hasattr(meta.creation_date, "isoformat")
                            else str(meta.creation_date)
                        )
                    if meta.modification_date:
                        document_metadata["modification_date"] = (
                            meta.modification_date.isoformat()
                            if hasattr(meta.modification_date, "isoformat")
                            else str(meta.modification_date)
                        )
                    # Keywords (may be comma-separated string)
                    if hasattr(meta, 'keywords') and meta.keywords:
                        document_metadata["keywords"] = str(meta.keywords)

                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)

                full_text = "\n\n".join(pages)

                document_metadata["num_pages"] = len(reader.pages)
                return {
                    "text": full_text,
                    "pages": len(reader.pages),
                    "document_metadata": document_metadata,
                }

            except Exception as e:
                # Add context to error
                raise Exception(f"PDF reading error: {str(e)}")

        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, extract)

    async def _extract_docx(self, path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX file.

        Args:
            path: Path to DOCX file

        Returns:
            dict with text, paragraph count, and document_metadata

        Raises:
            ImportError: If python-docx is not installed
            Exception: For DOCX reading errors
        """
        if not self._has_docx:
            raise ImportError(
                "python-docx library not installed. "
                "Install with: pip install python-docx"
            )

        from docx import Document

        # Run in executor to avoid blocking
        def extract():
            try:
                doc = Document(str(path))

                # Extract document metadata from core properties
                document_metadata = {}
                if doc.core_properties:
                    props = doc.core_properties
                    if props.author:
                        document_metadata["author"] = str(props.author)
                    if props.title:
                        document_metadata["title"] = str(props.title)
                    if props.subject:
                        document_metadata["subject"] = str(props.subject)
                    if props.keywords:
                        document_metadata["keywords"] = str(props.keywords)
                    if props.category:
                        document_metadata["category"] = str(props.category)
                    if props.comments:
                        document_metadata["comments"] = str(props.comments)
                    if props.last_modified_by:
                        document_metadata["last_modified_by"] = str(props.last_modified_by)
                    if props.created:
                        document_metadata["creation_date"] = str(props.created)
                    if props.modified:
                        document_metadata["modification_date"] = str(props.modified)
                    if props.revision:
                        document_metadata["revision"] = str(props.revision)

                # Extract paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)

                # Extract table content
                tables = []
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):  # Skip empty rows
                            tables.append(" | ".join(cells))

                # Combine all text
                all_text = []
                all_text.extend(paragraphs)
                if tables:
                    all_text.append("\n--- Tables ---\n")
                    all_text.extend(tables)

                full_text = "\n".join(all_text)

                return {
                    "text": full_text,
                    "pages": len(paragraphs),  # Use paragraph count as proxy
                    "document_metadata": document_metadata,
                }

            except Exception as e:
                raise Exception(f"DOCX reading error: {str(e)}")

        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, extract)

    async def _extract_xlsx(self, path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from XLSX file.

        Args:
            path: Path to XLSX file

        Returns:
            dict with text, sheet count, and document_metadata

        Raises:
            ImportError: If openpyxl is not installed
            Exception: For XLSX reading errors
        """
        if not self._has_openpyxl:
            raise ImportError(
                "openpyxl library not installed. "
                "Install with: pip install openpyxl"
            )

        from openpyxl import load_workbook

        # Run in executor to avoid blocking
        def extract():
            try:
                # Load workbook (NOT read-only so we can access properties)
                wb = load_workbook(
                    str(path),
                    read_only=False,  # Need full load for properties
                    data_only=True  # Get computed values, not formulas
                )

                # Extract workbook metadata
                document_metadata = {}
                if wb.properties:
                    props = wb.properties
                    if props.creator:
                        document_metadata["author"] = str(props.creator)
                    if props.title:
                        document_metadata["title"] = str(props.title)
                    if props.subject:
                        document_metadata["subject"] = str(props.subject)
                    if props.description:
                        document_metadata["description"] = str(props.description)
                    if props.keywords:
                        document_metadata["keywords"] = str(props.keywords)
                    if props.category:
                        document_metadata["category"] = str(props.category)
                    if props.lastModifiedBy:
                        document_metadata["last_modified_by"] = str(props.lastModifiedBy)
                    if props.created:
                        document_metadata["creation_date"] = (
                            props.created.isoformat()
                            if hasattr(props.created, "isoformat")
                            else str(props.created)
                        )
                    if props.modified:
                        document_metadata["modification_date"] = (
                            props.modified.isoformat()
                            if hasattr(props.modified, "isoformat")
                            else str(props.modified)
                        )
                    if props.company:
                        document_metadata["company"] = str(props.company)

                sheets = []

                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]

                    sheet_text = [f"--- Sheet: {sheet_name} ---"]

                    # Extract cell values row by row
                    for row in sheet.iter_rows(values_only=True):
                        # Convert to strings and filter out None
                        cells = [str(cell) for cell in row if cell is not None]
                        if cells:
                            sheet_text.append(" | ".join(cells))

                    sheets.append("\n".join(sheet_text))

                full_text = "\n\n".join(sheets)

                return {
                    "text": full_text,
                    "pages": len(wb.sheetnames),
                    "document_metadata": document_metadata,
                }

            except Exception as e:
                raise Exception(f"XLSX reading error: {str(e)}")

        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, extract)

    async def _extract_text(self, path: Path) -> Dict[str, Any]:
        """
        Extract text from plain text file.

        Args:
            path: Path to text file

        Returns:
            dict with text and line count
        """
        def extract():
            try:
                # Try different encodings
                encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
                text = None

                for encoding in encodings:
                    try:
                        with open(path, "r", encoding=encoding) as f:
                            text = f.read()
                        break
                    except UnicodeDecodeError:
                        continue

                if text is None:
                    # Fallback: read as bytes and decode with errors ignored
                    with open(path, "rb") as f:
                        text = f.read().decode("utf-8", errors="replace")

                lines = text.count("\n") + 1

                return {
                    "text": text,
                    "pages": lines,  # Use line count as proxy for pages
                    "document_metadata": {},
                }

            except Exception as e:
                raise Exception(f"Text file reading error: {str(e)}")

        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, extract)

    async def _extract_csv(self, path: Path) -> Dict[str, Any]:
        """
        Extract text from CSV or TSV file.

        Converts tabular data to readable text format with headers and rows.

        Args:
            path: Path to CSV/TSV file

        Returns:
            dict with text, row count, and column info
        """
        import csv

        def extract():
            try:
                # Detect delimiter (CSV vs TSV)
                delimiter = "	" if path.suffix.lower() == ".tsv" else ","

                # Try different encodings
                encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
                rows = None

                for encoding in encodings:
                    try:
                        with open(path, "r", encoding=encoding, newline="") as f:
                            # Sniff to detect actual delimiter if not TSV
                            if path.suffix.lower() != ".tsv":
                                sample = f.read(4096)
                                f.seek(0)
                                try:
                                    dialect = csv.Sniffer().sniff(sample, delimiters=",;	|")
                                    delimiter = dialect.delimiter
                                except csv.Error:
                                    delimiter = ","  # Default to comma

                            reader = csv.reader(f, delimiter=delimiter)
                            rows = list(reader)
                        break
                    except UnicodeDecodeError:
                        continue

                if rows is None:
                    raise Exception("Could not decode CSV file with any supported encoding")

                if not rows:
                    return {
                        "text": "",
                        "pages": 0,
                        "document_metadata": {"columns": 0, "rows": 0},
                    }

                # Format as readable text
                # First row is typically headers
                headers = rows[0] if rows else []
                data_rows = rows[1:] if len(rows) > 1 else []

                text_parts = []

                # Add header line
                if headers:
                    text_parts.append("--- Columns ---")
                    text_parts.append(" | ".join(str(h) for h in headers))
                    text_parts.append("")
                    text_parts.append("--- Data ---")

                # Add data rows
                for row in data_rows:
                    # Format each row, optionally with column names
                    if headers and len(row) == len(headers):
                        # Format as "Column: Value" pairs for better readability
                        pairs = [f"{headers[i]}: {row[i]}" for i in range(len(row))]
                        text_parts.append(" | ".join(pairs))
                    else:
                        text_parts.append(" | ".join(str(cell) for cell in row))

                return {
                    "text": "\n".join(text_parts),
                    "pages": len(data_rows),  # Use row count as page proxy
                    "document_metadata": {
                        "columns": len(headers),
                        "rows": len(data_rows),
                        "headers": headers[:20],  # First 20 column names
                    },
                }

            except Exception as e:
                raise Exception(f"CSV reading error: {str(e)}")

        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, extract)

    async def _extract_eml(self, path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from EML or EMLX (Apple Mail) email file.

        Args:
            path: Path to email file

        Returns:
            dict with text (headers + body), part count, and document_metadata
        """
        import email
        from email import policy

        def extract():
            try:
                # Read file content
                with open(path, "rb") as f:
                    content = f.read()

                # Handle EMLX format (Apple Mail)
                # EMLX has a preamble with byte count on first line
                if path.suffix.lower() == ".emlx":
                    # Skip the preamble (first line is byte count)
                    lines = content.split(b"\n", 1)
                    if len(lines) > 1:
                        try:
                            # Check if first line is a number (byte count)
                            int(lines[0].strip())
                            content = lines[1]
                        except ValueError:
                            pass  # Not EMLX format, use as-is

                # Parse email
                msg = email.message_from_bytes(content, policy=policy.default)

                # Extract email metadata (headers are the metadata for emails!)
                document_metadata = {}

                # Core email metadata
                if msg.get("From"):
                    document_metadata["author"] = str(msg.get("From"))
                    document_metadata["email_from"] = str(msg.get("From"))
                if msg.get("To"):
                    document_metadata["email_to"] = str(msg.get("To"))
                if msg.get("Cc"):
                    document_metadata["email_cc"] = str(msg.get("Cc"))
                if msg.get("Bcc"):
                    document_metadata["email_bcc"] = str(msg.get("Bcc"))
                if msg.get("Subject"):
                    document_metadata["title"] = str(msg.get("Subject"))
                    document_metadata["email_subject"] = str(msg.get("Subject"))
                if msg.get("Date"):
                    document_metadata["creation_date"] = str(msg.get("Date"))
                    document_metadata["email_date"] = str(msg.get("Date"))

                # Additional email metadata
                if msg.get("Reply-To"):
                    document_metadata["email_reply_to"] = str(msg.get("Reply-To"))
                if msg.get("Message-ID"):
                    document_metadata["email_message_id"] = str(msg.get("Message-ID"))
                if msg.get("In-Reply-To"):
                    document_metadata["email_in_reply_to"] = str(msg.get("In-Reply-To"))
                if msg.get("X-Mailer"):
                    document_metadata["creator"] = str(msg.get("X-Mailer"))
                if msg.get("User-Agent"):
                    document_metadata["creator"] = str(msg.get("User-Agent"))
                if msg.get("Organization"):
                    document_metadata["organization"] = str(msg.get("Organization"))

                # Check for attachments
                attachment_count = 0
                attachment_names = []
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_disposition() == "attachment":
                            attachment_count += 1
                            filename = part.get_filename()
                            if filename:
                                attachment_names.append(filename)

                if attachment_count > 0:
                    document_metadata["attachment_count"] = attachment_count
                    if attachment_names:
                        document_metadata["attachments"] = ", ".join(attachment_names)

                parts = []
                part_count = 0

                # Extract headers for text content
                headers = []
                for header in ["From", "To", "Cc", "Subject", "Date"]:
                    value = msg.get(header)
                    if value:
                        headers.append(f"{header}: {value}")

                if headers:
                    parts.append("--- Headers ---")
                    parts.extend(headers)
                    parts.append("")

                # Extract body
                parts.append("--- Body ---")

                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        if content_type == "text/plain":
                            body = part.get_content()
                            if isinstance(body, str):
                                parts.append(body)
                                part_count += 1
                        elif content_type == "text/html":
                            # Basic HTML stripping
                            import re
                            html = part.get_content()
                            if isinstance(html, str):
                                text = re.sub(r"<[^>]+>", "", html)
                                text = re.sub(r"\s+", " ", text).strip()
                                if text and part_count == 0:  # Only if no plain text
                                    parts.append(text)
                                    part_count += 1
                else:
                    body = msg.get_content()
                    if isinstance(body, str):
                        parts.append(body)
                        part_count = 1
                    elif isinstance(body, bytes):
                        parts.append(body.decode("utf-8", errors="replace"))
                        part_count = 1

                return {
                    "text": "\n".join(parts),
                    "pages": max(1, part_count),
                    "document_metadata": document_metadata,
                }

            except Exception as e:
                raise Exception(f"Email file reading error: {str(e)}")

        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, extract)


if __name__ == "__main__":
    """Run the worker if executed directly."""
    from arkham_frame.workers.base import run_worker

    run_worker(ExtractWorker)
